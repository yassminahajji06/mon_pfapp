from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    Image as RLImage,
    KeepTogether,
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
TUTORIAL_DIR = DOCS / "tutorial"
PFE_DIR = DOCS / "pfe-report"
ASSET_DIR = TUTORIAL_DIR / "assets"
SOURCE_DIR = ASSET_DIR / "source"
SCREENSHOT_DIR = ASSET_DIR / "screenshots"
PFE_ASSET_DIR = PFE_DIR / "assets"
LOGO_PATH = SOURCE_DIR / "cover_image_1.png"

TUTORIAL_DOCX = TUTORIAL_DIR / "Mon_PF_App_Tutorial_Yassmine_Hajji.docx"
TUTORIAL_PDF = TUTORIAL_DIR / "Mon_PF_App_Tutorial_Yassmine_Hajji.pdf"
PFE_DOCX = PFE_DIR / "Rapport_PFE_Yassmine_Hajji_Mon_PF_App.docx"
PFE_PDF = PFE_DIR / "Rapport_PFE_Yassmine_Hajji_Mon_PF_App.pdf"

RED = RGBColor(229, 57, 53)
BLUE = RGBColor(46, 116, 181)
NAVY = RGBColor(31, 77, 120)
TEXT = RGBColor(31, 41, 51)
MUTED = RGBColor(93, 104, 119)
LIGHT_BLUE = "E8EEF5"
LIGHT_RED = "FDECEC"
LIGHT_GREEN = "EAF7ED"
LIGHT_YELLOW = "FFF7E6"
LIGHT_GRAY = "F3F4F6"


@dataclass(frozen=True)
class ProjectInfo:
    student: str = "HAJJI Yassmine"
    formation: str = "BTS Developpement des Applications Informatiques"
    school: str = "Lycee Qualifiant Technique Ibn Al-Haitam"
    ministry: str = "Royaume du Maroc - Ministere de l'Education Nationale, du Prescolaire et des Sports"
    academy: str = "Academie Regionale de l'Education et de la Formation Draa-Tafilalet"
    direction: str = "Direction Provinciale de Ouarzazate"
    academic_year: str = "2024 - 2026"
    project: str = "Mon PF App"
    report_title: str = "Conception et realisation d'une application mobile de gestion des commandes d'un restaurant"
    supervisor: str = "Encadrant pedagogique : a completer"


INFO = ProjectInfo()


def font_path(name: str) -> str | None:
    candidates = [
        Path("C:/Windows/Fonts") / name,
        Path("/usr/share/fonts/truetype/dejavu") / name,
    ]
    for candidate in candidates:
        if candidate.exists():
            return str(candidate)
    return None


def pil_font(size: int, bold: bool = False) -> ImageFont.ImageFont:
    names = ("arialbd.ttf", "DejaVuSans-Bold.ttf") if bold else ("arial.ttf", "DejaVuSans.ttf")
    for name in names:
        path = font_path(name)
        if path:
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def sanitize(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def set_run_font(
    run,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def paragraph_border_bottom(paragraph, color: str = "D7DBE2", size: str = "8") -> None:
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = "DADCE0") -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def setup_doc(doc: Document, *, compact: bool = False) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.75 if compact else 0.82)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5 if compact else 11)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(6 if compact else 8)
    normal.paragraph_format.line_spacing = 1.18 if compact else 1.25

    for level, size, color, before, after in [
        (1, 16, BLUE, 16, 8),
        (2, 13, BLUE, 12, 6),
        (3, 12, NAVY, 8, 4),
    ]:
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = ""
    run = header.add_run(f"{INFO.project} | {INFO.student} | BTS DAI")
    set_run_font(run, size=8.5, color=MUTED, italic=True)
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    footer = section.footer.paragraphs[0]
    footer.text = "Rapport PFE - Projet de fin d'etudes"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(footer.runs[0], size=8.5, color=MUTED)


def add_para(
    doc: Document,
    text: str,
    *,
    bold_lead: str | None = None,
    style: str | None = None,
    align: WD_ALIGN_PARAGRAPH | None = None,
) -> None:
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, bold=True)
        rest = p.add_run(text[len(bold_lead) :])
        set_run_font(rest)
    else:
        run = p.add_run(text)
        set_run_font(run)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_run_font(run)


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_run_font(run)


def add_code(doc: Document, code: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color="E5E7EB")
    cell = table.cell(0, 0)
    cell.width = Inches(6.4)
    shade_cell(cell, "111827")
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    for i, line in enumerate(code.splitlines()):
        if i:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run_font(run, name="Consolas", size=8.6, color=RGBColor(245, 245, 245))
    doc.add_paragraph()


def add_matrix(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color="CBD5E1")
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        cell.width = Inches(widths[idx])
        shade_cell(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_run_font(run, size=9.2, color=NAVY, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            cell.width = Inches(widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if len(value) > 18 else WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(value)
            set_run_font(run, size=8.8)
    doc.add_paragraph()


def add_image(doc: Document, path: Path, caption: str, *, width: float = 5.8) -> None:
    if not path.exists():
        add_para(doc, f"[Image manquante: {path.name}]", align=WD_ALIGN_PARAGRAPH.CENTER)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = c.add_run(caption)
    set_run_font(run, size=8.8, color=MUTED, italic=True)


def add_report_cover(doc: Document) -> None:
    for line in [INFO.ministry, INFO.academy, INFO.direction]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        set_run_font(r, size=9.5, color=NAVY, bold=True)
    if LOGO_PATH.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO_PATH), width=Inches(1.55))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Rapport de Projet de Fin d'Etudes")
    set_run_font(r, size=23, color=TEXT, bold=True)
    paragraph_border_bottom(p, color="E53935", size="12")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Deuxieme annee BTS")
    set_run_font(r, size=13, color=MUTED, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(INFO.formation)
    set_run_font(r, size=13, color=NAVY, bold=True)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(INFO.report_title.upper())
    set_run_font(r, size=16.5, color=RED, bold=True)

    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color="E5E7EB")
    cover_rows = [
        ("Projet", INFO.project),
        ("Realise par", INFO.student),
        ("Etablissement", INFO.school),
        ("Encadrement", INFO.supervisor),
        ("Annee de formation", INFO.academic_year),
    ]
    for row_idx, (label, value) in enumerate(cover_rows):
        cells = table.rows[row_idx].cells
        for idx, txt in enumerate((label, value)):
            cells[idx].width = Inches(2.0 if idx == 0 else 4.1)
            set_cell_margins(cells[idx], top=120, bottom=120, start=140, end=140)
            shade_cell(cells[idx], LIGHT_GRAY if idx == 0 else "FFFFFF")
            run = cells[idx].paragraphs[0].add_run(txt)
            set_run_font(run, size=10, color=NAVY if idx == 0 else TEXT, bold=idx == 0)
    doc.add_page_break()


def wrap_text(text: str, font: ImageFont.ImageFont, max_width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        test = f"{current} {word}".strip()
        if font.getbbox(test)[2] <= max_width or not current:
            current = test
        else:
            lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_box(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], title: str, lines: list[str], fill: str) -> None:
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=22, fill=fill, outline="#CBD5E1", width=2)
    draw.text((x1 + 20, y1 + 18), title, fill="#1F4D78", font=pil_font(24, True))
    y = y1 + 58
    font = pil_font(18)
    for line in lines:
        for wrapped in wrap_text(line, font, x2 - x1 - 40):
            draw.text((x1 + 20, y), wrapped, fill="#374151", font=font)
            y += 24


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = "#E53935") -> None:
    draw.line([start, end], fill=color, width=5)
    x, y = end
    draw.polygon([(x, y), (x - 18, y - 9), (x - 18, y + 9)], fill=color)


def build_diagrams() -> dict[str, Path]:
    PFE_ASSET_DIR.mkdir(parents=True, exist_ok=True)
    diagrams: dict[str, Path] = {}

    canvas = Image.new("RGB", (1500, 900), "white")
    draw = ImageDraw.Draw(canvas)
    draw.rectangle((0, 0, 1500, 90), fill="#E53935")
    draw.text((40, 25), "Architecture demo-prod", fill="white", font=pil_font(34, True))
    draw_box(draw, (55, 160, 345, 345), "Flutter", ["UI client", "Panier", "Admin", "Livreur"], "#FDECEC")
    draw_box(draw, (435, 160, 725, 345), "Services API", ["ApiClient", "MenuService", "OrderService", "AuthService"], "#E8EEF5")
    draw_box(draw, (815, 160, 1105, 345), "Laravel API", ["Routes REST", "Validations", "Roles serveur", "Token Bearer"], "#EAF7ED")
    draw_box(draw, (1195, 160, 1440, 345), "SQLite", ["users", "menu_items", "orders", "order_items"], "#FFF7E6")
    draw_box(draw, (435, 560, 725, 725), "Secure Storage", ["Token", "Role", "Suppression logout"], "#F3F4F6")
    draw_box(draw, (815, 560, 1105, 725), "Tests", ["PHPUnit", "Flutter tests", "Dart analyze", "APK debug"], "#F3F4F6")
    for s, e in [((345, 252), (435, 252)), ((725, 252), (815, 252)), ((1105, 252), (1195, 252))]:
        draw_arrow(draw, s, e)
    draw_arrow(draw, (580, 345), (580, 560), "#43A047")
    draw_arrow(draw, (960, 345), (960, 560), "#1E88E5")
    diagrams["architecture"] = PFE_ASSET_DIR / "architecture_demo_prod.png"
    canvas.save(diagrams["architecture"])

    canvas = Image.new("RGB", (1500, 900), "white")
    draw = ImageDraw.Draw(canvas)
    draw.rectangle((0, 0, 1500, 90), fill="#1E88E5")
    draw.text((40, 25), "Cycle de commande", fill="white", font=pil_font(34, True))
    steps = [
        ("Client", "Ajoute plats au panier", "#FDECEC"),
        ("Commande", "POST /api/orders", "#E8EEF5"),
        ("Restaurant", "Preparation / Pret", "#EAF7ED"),
        ("Livreur", "Accepte livraison", "#FFF7E6"),
        ("Client", "Suit le statut", "#F3F4F6"),
    ]
    x = 55
    for title, body, fill in steps:
        draw_box(draw, (x, 250, x + 245, 430), title, [body], fill)
        if x < 1100:
            draw_arrow(draw, (x + 245, 340), (x + 330, 340))
        x += 300
    draw.text((70, 650), "Regle: une commande ne peut pas etre creee sans plats. Les statuts sont controles par l'API.", fill="#374151", font=pil_font(26, True))
    diagrams["order_cycle"] = PFE_ASSET_DIR / "order_cycle.png"
    canvas.save(diagrams["order_cycle"])

    canvas = Image.new("RGB", (1500, 900), "white")
    draw = ImageDraw.Draw(canvas)
    draw.rectangle((0, 0, 1500, 90), fill="#43A047")
    draw.text((40, 25), "Modele logique simplifie", fill="white", font=pil_font(34, True))
    boxes = [
        ((70, 160, 390, 330), "users", ["id, name, email", "role, password", "api_token_hash"]),
        ((590, 160, 910, 330), "orders", ["client, driver", "status, address", "subtotal, total"]),
        ((1110, 160, 1430, 330), "order_items", ["order_id", "menu_item_id", "quantity, price"]),
        ((590, 560, 910, 730), "menu_items", ["category_id", "name, price", "available"]),
        ((1110, 560, 1430, 730), "categories", ["slug, label", "color, icon"]),
    ]
    for xy, title, lines in boxes:
        draw_box(draw, xy, title, lines, "#F3F4F6")
    draw_arrow(draw, (390, 245), (590, 245), "#E53935")
    draw_arrow(draw, (910, 245), (1110, 245), "#E53935")
    draw_arrow(draw, (1270, 330), (750, 560), "#1E88E5")
    draw_arrow(draw, (1110, 640), (910, 640), "#1E88E5")
    diagrams["data_model"] = PFE_ASSET_DIR / "data_model_backend.png"
    canvas.save(diagrams["data_model"])

    return diagrams


REPORT_SECTIONS = [
    ("Remerciements", [
        "Je tiens a remercier l'equipe pedagogique du BTS Developpement des Applications Informatiques du Lycee Qualifiant Technique Ibn Al-Haitam pour l'encadrement assure durant ma formation.",
        "Je remercie egalement toutes les personnes qui ont contribue, directement ou indirectement, a la definition du besoin, aux tests et a l'amelioration du projet Mon PF App.",
    ]),
    ("Resume", [
        "Ce rapport presente la conception et la realisation de Mon PF App, une application mobile de gestion des commandes et des livraisons d'un restaurant. Le projet repond a une problematique courante: la gestion manuelle provoque des erreurs, des retards et une mauvaise communication entre client, restaurant et livreur.",
        "La solution proposee s'appuie sur une application Flutter et une API Laravel. Elle permet au client de consulter le menu, gerer le panier, passer une commande et suivre son etat. Elle propose aussi un espace administrateur et un espace livreur. Pour la soutenance, le projet dispose d'un mode demonstration fiable et d'un backend local demo-prod avec base SQLite.",
    ]),
]


def add_static_toc(doc: Document, title: str, items: list[str]) -> None:
    doc.add_heading(title, level=1)
    for idx, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        run = p.add_run(f"{idx}. {item}")
        set_run_font(run, size=10.2)
    doc.add_page_break()


def build_pfe_docx(diagrams: dict[str, Path]) -> None:
    PFE_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_doc(doc, compact=False)
    add_report_cover(doc)

    for heading, paragraphs in REPORT_SECTIONS:
        doc.add_heading(heading, level=1)
        for paragraph in paragraphs:
            add_para(doc, paragraph)
        doc.add_page_break()

    toc_items = [
        "Introduction generale",
        "Contexte et problematique",
        "Cahier des charges",
        "Analyse et conception",
        "Architecture technique",
        "Realisation backend",
        "Realisation Flutter",
        "Securite et qualite",
        "Tests et validation",
        "Installation, build et mise en service",
        "Difficultes rencontrees",
        "Conclusion et perspectives",
        "Annexes",
    ]
    add_static_toc(doc, "Table des matieres", toc_items)

    doc.add_heading("1. Introduction generale", level=1)
    add_para(doc, "Dans le cadre de la formation BTS Developpement des Applications Informatiques, ce projet de fin d'etudes consiste a concevoir et realiser une application mobile de gestion des commandes d'un restaurant.")
    add_para(doc, "Le projet Mon PF App permet de relier plusieurs acteurs: le client, l'administrateur, le serveur et le livreur. Il couvre un parcours complet depuis la consultation du menu jusqu'au suivi de livraison.")
    add_para(doc, "Le travail realise ne se limite pas a une interface graphique. Il inclut une architecture frontend Flutter, un backend Laravel, une base de donnees locale SQLite, des tests automatiques, une documentation technique et un guide de mise en service.")
    add_matrix(doc, ["Element", "Description"], [
        ["Sujet", INFO.report_title],
        ["Stagiaire", INFO.student],
        ["Formation", INFO.formation],
        ["Etablissement", INFO.school],
        ["Projet", INFO.project],
    ], [1.8, 4.7])

    doc.add_heading("2. Contexte et problematique", level=1)
    add_para(doc, "Les restaurants font face a une augmentation des commandes, en particulier les commandes a emporter et les livraisons. Lorsque la gestion reste manuelle, plusieurs problemes apparaissent: erreurs de commande, retards, mauvaise communication et insatisfaction des clients.")
    add_para(doc, "L'application proposee digitalise ce processus. Elle centralise les commandes, automatise le calcul du total, affiche l'etat de suivi et separe les responsabilites selon les roles.")
    add_matrix(doc, ["Probleme", "Impact", "Reponse du projet"], [
        ["Erreur de commande", "Perte de temps et frustration client", "Panier numerique et validation explicite"],
        ["Retard de preparation", "Mauvaise experience utilisateur", "Statuts de commande suivis"],
        ["Communication faible", "Client/livreur/restaurant mal synchronises", "Espaces dedies par acteur"],
        ["Gestion difficile du menu", "Prix et disponibilites non maitrises", "Endpoints admin de gestion menu"],
    ], [1.7, 2.2, 2.6])

    doc.add_heading("3. Cahier des charges", level=1)
    add_para(doc, "Le cahier de charge demande une application mobile permettant d'organiser les commandes, suivre les plats, ameliorer la rapidite du service et reduire les erreurs humaines.")
    add_matrix(doc, ["Acteur", "Besoins fonctionnels couverts"], [
        ["Client", "Creation compte, connexion, consultation menu, ajout panier, passage commande, suivi de commande"],
        ["Administrateur", "Statistiques, commandes, gestion menu cote API, affectation livreur cote API"],
        ["Serveur", "Role prevu cote API; flux operationnel rattache a l'administration"],
        ["Livreur", "Connexion, consultation file livraison, acceptation livraison, suivi statut"],
    ], [1.45, 5.05])
    add_bullets(doc, [
        "Securite des donnees et authentification par token.",
        "Interface mobile simple, visuelle et adaptee au domaine restaurant.",
        "Rapidite de traitement grace aux appels API et au mode demo de secours.",
        "Maintenance facilitee par une structure de projet separee frontend/backend.",
    ])

    doc.add_heading("4. Analyse et conception", level=1)
    add_para(doc, "L'analyse distingue quatre acteurs principaux: Client, Administrateur, Serveur et Livreur. Le systeme est organise autour des entites Utilisateur, Plat, Panier, Commande et Livraison.")
    add_image(doc, diagrams["data_model"], "Figure 1 - Modele logique simplifie du backend.", width=6.3)
    add_matrix(doc, ["Entite", "Role dans le systeme"], [
        ["User", "Stocke les informations utilisateur, le role et le token hash."],
        ["MenuItem", "Represente un plat avec prix, categorie, disponibilite et image."],
        ["Order", "Represente une commande client avec statut, adresse, total et livreur affecte."],
        ["OrderItem", "Stocke les lignes de commande afin de conserver prix et quantite."],
        ["Category", "Organise les plats du menu par familles."],
    ], [1.55, 4.95])

    doc.add_heading("5. Architecture technique", level=1)
    add_image(doc, diagrams["architecture"], "Figure 2 - Architecture demo-prod Flutter + Laravel.", width=6.4)
    add_para(doc, "L'architecture separe clairement le frontend mobile et le backend. Flutter gere l'experience utilisateur et l'etat local du panier. Laravel expose une API REST et garantit les validations metier. SQLite sert de base locale de demonstration, facile a lancer sans serveur externe.")
    add_matrix(doc, ["Couche", "Technologie", "Justification"], [
        ["Frontend", "Flutter / Dart", "Un seul code pour Android, Windows et web de test."],
        ["Backend", "Laravel 13", "Framework structure, routes REST, validations, tests."],
        ["Base demo", "SQLite", "Demarrage rapide pour soutenance sans installation MySQL."],
        ["Securite locale", "flutter_secure_storage", "Conservation plus sure du token."],
        ["Tests", "PHPUnit + flutter_test", "Verification automatique des flux principaux."],
    ], [1.2, 1.7, 3.6])

    doc.add_heading("6. Realisation backend", level=1)
    add_para(doc, "Le backend se trouve dans le dossier mon_pfapi. Il contient les routes API, les controleurs, les modeles Eloquent, les migrations et les seeders.")
    add_matrix(doc, ["Fichier", "Responsabilite"], [
        ["routes/api.php", "Expose health, menu, register, login, orders, admin stats, driver queue."],
        ["AuthController.php", "Inscription, connexion, deconnexion et emission token."],
        ["MenuController.php", "Catalogue public et endpoints admin de gestion menu."],
        ["OrderController.php", "Creation commande, historique, statuts, affectation et file livreur."],
        ["TokenAuth.php", "Middleware Bearer token + controle roles."],
        ["DatabaseSeeder.php", "Comptes test, restaurant, categories, menu et commande initiale."],
    ], [2.1, 4.4])
    add_para(doc, "Les roles sensibles ne sont jamais choisis par l'utilisateur public. Lors de l'inscription, le backend force le role client. Les comptes admin, serveur et livreur sont crees cote seed ou administration.")

    doc.add_heading("7. Realisation Flutter", level=1)
    add_para(doc, "Le frontend Flutter se trouve dans mon_pfapp. Il est organise par domaines: auth, client, operations, modeles et widgets partages.")
    add_matrix(doc, ["Dossier/Fichier", "Role"], [
        ["lib/app/mon_pf_app.dart", "Etat global, navigation, panier, chargement menu et commandes."],
        ["lib/data/api_client.dart", "Client HTTP avec headers, token, timeout et erreurs."],
        ["features/client/data", "MenuService et OrderService."],
        ["features/auth/data/auth_service.dart", "Login, register, logout."],
        ["shared/widgets/app_ui.dart", "Couleurs, cartes, navigation, visuels de plats."],
    ], [2.4, 4.1])
    add_image(doc, SCREENSHOT_DIR / "contact-sheet.png", "Figure 3 - Apercu des principaux ecrans Flutter.", width=6.3)

    doc.add_heading("8. Securite et qualite", level=1)
    add_para(doc, "Plusieurs decisions rapprochent le projet d'un comportement production: configuration API par environnement, stockage securise du token, suppression du choix de role public, timeouts reseau et checks mounted dans les ecrans asynchrones.")
    add_matrix(doc, ["Risque", "Solution"], [
        ["Auto-creation admin", "Role force cote serveur."],
        ["Token expose", "flutter_secure_storage cote app et hash du token cote API."],
        ["Serveur indisponible", "Timeout et messages d'erreur comprehensibles."],
        ["HTTP en release", "Release Android reste stricte; HTTP autorise seulement debug/profile."],
        ["Regression UI", "Tests widget sur login, inscription, menu, panier, suivi, espaces admin/livreur."],
    ], [2.1, 4.4])

    doc.add_heading("9. Tests et validation", level=1)
    add_matrix(doc, ["Test", "Commande", "Resultat"], [
        ["Analyse Flutter", "dart.exe analyze", "No issues found"],
        ["Tests Flutter", "flutter.bat test", "5 tests passes"],
        ["Tests Laravel", "php artisan test", "5 tests passes"],
        ["Flux API live", "login + POST /orders", "Commande #PF-0851 creee"],
        ["Android APK debug", "flutter build apk --debug", "APK genere"],
        ["Windows", "flutter build windows --debug", "Bloque par composant ATL manquant"],
    ], [1.6, 3.1, 1.8])
    add_image(doc, diagrams["order_cycle"], "Figure 4 - Cycle de validation d'une commande.", width=6.4)

    doc.add_heading("10. Installation, build et mise en service", level=1)
    add_para(doc, "Le projet peut etre lance en mode demo ou en mode backend reel. Les commandes exactes sont documentees pour eviter les erreurs de chemin.")
    add_code(doc, r'''cd "C:\Users\LOQ\Documents\yassmin pfe\mon_pfapi"
..\tools\runtime\php-8.4.22\php.exe artisan migrate:fresh --seed
..\tools\runtime\php-8.4.22\php.exe artisan serve --host=0.0.0.0 --port=8000''')
    add_code(doc, r'''cd "C:\Users\LOQ\Documents\yassmin pfe\mon_pfapp"
C:\Users\LOQ\dev\flutter\bin\flutter.bat run --dart-define=DEMO_MODE=false --dart-define=API_BASE_URL=http://127.0.0.1:8000/api''')
    add_para(doc, "Pour telephone Android sur le meme Wi-Fi, remplacer 127.0.0.1 par l'adresse IP Wi-Fi du PC. L'APK debug actuel pointe vers http://172.17.182.162:8000/api.")

    doc.add_heading("11. Difficultes rencontrees", level=1)
    add_matrix(doc, ["Probleme", "Diagnostic", "Solution"], [
        ["Installation PHP globale", "Chocolatey bloque par permissions ProgramData", "Installation portable PHP/Composer dans tools/runtime."],
        ["dart.bat lent/bloque", "Wrapper Windows problematique", "Utiliser dart.exe direct du SDK Flutter."],
        ["Android HTTP local", "Release refuse cleartext", "Autorisation HTTP seulement debug/profile avec tools:replace."],
        ["Windows build", "atlstr.h manquant", "Installer Visual Studio C++ ATL."],
        ["Telephone non detecte", "ADB pas dans PATH ou USB debug absent", "Utiliser adb.exe par chemin complet et accepter RSA."],
    ], [1.7, 2.4, 2.4])

    doc.add_heading("12. Conclusion et perspectives", level=1)
    add_para(doc, "Mon PF App valide une grande partie du cahier de charge: parcours client, menu, panier, commande, suivi, administration, espace livreur, backend API et tests. Le projet est presentable comme produit demo-prod pour une soutenance BTS.")
    add_para(doc, "Les perspectives concernent surtout la production reelle: hebergement HTTPS, base MySQL/PostgreSQL, Laravel Sanctum, notifications push, paiement en ligne, synchronisation temps reel et interface admin complete pour le CRUD menu.")

    doc.add_heading("Annexes", level=1)
    add_matrix(doc, ["Compte", "Role", "Mot de passe"], [
        ["yassmine@monpf.fr", "client", "password"],
        ["admin@monpf.fr", "admin", "password"],
        ["serveur@monpf.fr", "serveur", "password"],
        ["livreur@monpf.fr", "livreur", "password"],
    ], [2.5, 1.7, 2.3])
    add_matrix(doc, ["Document", "Emplacement"], [
        ["Tutoriel technique", "docs/tutorial/Mon_PF_App_Tutorial_Yassmine_Hajji.pdf"],
        ["Runbook demo-prod", "docs/demo-prod-runbook.md"],
        ["API backend", "docs/backend-api.md"],
        ["Checklist validation", "docs/specifications/validation-checklist.md"],
        ["APK debug", "mon_pfapp/build/app/outputs/flutter-apk/app-debug.apk"],
    ], [2.2, 4.3])
    doc.save(PFE_DOCX)


def add_tutorial_cover(doc: Document) -> None:
    if LOGO_PATH.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO_PATH), width=Inches(1.35))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Mon PF App")
    set_run_font(r, size=26, color=RED, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Tutoriel complet d'installation, test, debug, build et demonstration")
    set_run_font(r, size=14, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{INFO.student} - {INFO.formation}")
    set_run_font(r, size=11, color=MUTED)
    doc.add_page_break()


def build_tutorial_docx(diagrams: dict[str, Path]) -> None:
    doc = Document()
    setup_doc(doc, compact=True)
    add_tutorial_cover(doc)
    add_static_toc(doc, "Sommaire du tutoriel", [
        "Ce que contient le projet",
        "Pre-requis et chemins importants",
        "Lancer le backend Laravel",
        "Lancer Flutter en mode demo ou backend reel",
        "Tester sur Android",
        "Tester sur Windows",
        "Commandes de test et de build",
        "Guide de debug",
        "Architecture et fichiers a comprendre",
        "Scenario de soutenance",
    ])

    doc.add_heading("1. Ce que contient le projet", level=1)
    add_para(doc, "Le repo contient deux applications complementaires: mon_pfapp pour Flutter et mon_pfapi pour Laravel. Le premier affiche l'interface mobile, le second fournit les donnees reelles demo-prod.")
    add_matrix(doc, ["Dossier", "Contenu"], [
        ["mon_pfapp", "Application Flutter Android/Windows/web de test."],
        ["mon_pfapi", "Backend Laravel 13 avec API REST et SQLite."],
        ["docs", "Cahier de charge, rapport PFE, tutoriel, runbooks et design."],
        ["tools", "Scripts de generation documents/captures et runtime PHP portable ignore par Git."],
    ], [1.6, 4.9])

    doc.add_heading("2. Pre-requis et chemins importants", level=1)
    add_matrix(doc, ["Outil", "Chemin / remarque"], [
        ["Flutter", r"C:\Users\LOQ\dev\flutter\bin\flutter.bat"],
        ["Dart direct", r"C:\Users\LOQ\dev\flutter\bin\cache\dart-sdk\bin\dart.exe"],
        ["PHP portable", r"C:\Users\LOQ\Documents\yassmin pfe\tools\runtime\php-8.4.22\php.exe"],
        ["ADB", r"C:\Users\LOQ\AppData\Local\Android\Sdk\platform-tools\adb.exe"],
        ["Repo", r"C:\Users\LOQ\Documents\yassmin pfe"],
    ], [1.5, 5.0])
    add_para(doc, "Important: sur cette machine, utiliser dart.exe direct pour l'analyse. Le wrapper dart.bat peut prendre trop de temps.")

    doc.add_heading("3. Lancer le backend Laravel", level=1)
    add_code(doc, r'''cd "C:\Users\LOQ\Documents\yassmin pfe\mon_pfapi"
..\tools\runtime\php-8.4.22\php.exe artisan migrate:fresh --seed
..\tools\runtime\php-8.4.22\php.exe artisan serve --host=0.0.0.0 --port=8000''')
    add_para(doc, "Tester que l'API repond:")
    add_code(doc, r'''Invoke-RestMethod http://127.0.0.1:8000/api/health
Invoke-RestMethod http://127.0.0.1:8000/api/menu''')
    add_matrix(doc, ["Compte", "Role", "Mot de passe"], [
        ["yassmine@monpf.fr", "client", "password"],
        ["admin@monpf.fr", "admin", "password"],
        ["serveur@monpf.fr", "serveur", "password"],
        ["livreur@monpf.fr", "livreur", "password"],
    ], [2.4, 1.7, 2.2])

    doc.add_heading("4. Lancer Flutter", level=1)
    add_para(doc, "Mode demo sans backend:")
    add_code(doc, r'''cd "C:\Users\LOQ\Documents\yassmin pfe\mon_pfapp"
C:\Users\LOQ\dev\flutter\bin\flutter.bat run''')
    add_para(doc, "Mode backend reel sur Windows:")
    add_code(doc, r'''C:\Users\LOQ\dev\flutter\bin\flutter.bat run -d windows --dart-define=DEMO_MODE=false --dart-define=API_BASE_URL=http://127.0.0.1:8000/api''')
    add_para(doc, "Mode backend reel sur telephone Android Wi-Fi:")
    add_code(doc, r'''C:\Users\LOQ\dev\flutter\bin\flutter.bat run -d <device-id> --dart-define=DEMO_MODE=false --dart-define=API_BASE_URL=http://ADRESSE_IP_PC:8000/api''')

    doc.add_heading("5. Tester sur Android", level=1)
    add_numbered(doc, [
        "Activer Options developpeur sur le telephone.",
        "Activer Debogage USB.",
        "Brancher le telephone en USB.",
        "Accepter la popup RSA sur le telephone.",
        "Verifier avec adb devices.",
        "Installer l'APK debug ou lancer flutter run.",
    ])
    add_code(doc, r'''C:\Users\LOQ\AppData\Local\Android\Sdk\platform-tools\adb.exe devices
C:\Users\LOQ\AppData\Local\Android\Sdk\platform-tools\adb.exe install -r build\app\outputs\flutter-apk\app-debug.apk''')
    add_para(doc, "APK debug actuel:")
    add_code(doc, r'''C:\Users\LOQ\Documents\yassmin pfe\mon_pfapp\build\app\outputs\flutter-apk\app-debug.apk''')
    add_para(doc, "Si le telephone ne voit pas l'API, verifier que le telephone et le PC sont sur le meme Wi-Fi, que le backend tourne, et que l'URL contient l'IPv4 Wi-Fi du PC.")

    doc.add_heading("6. Tester sur Windows", level=1)
    add_para(doc, "Le build Windows est actuellement bloque par le fichier atlstr.h manquant. Ce fichier vient du composant Visual Studio ATL.")
    add_numbered(doc, [
        "Ouvrir Visual Studio Installer.",
        "Modifier Visual Studio Community 2026.",
        "Chercher ATL dans Individual components.",
        "Installer C++ ATL for latest v143/v14x build tools ou equivalent.",
        "Relancer flutter build windows.",
    ])
    add_code(doc, r'''C:\Users\LOQ\dev\flutter\bin\flutter.bat build windows --debug --dart-define=DEMO_MODE=false --dart-define=API_BASE_URL=http://127.0.0.1:8000/api''')

    doc.add_heading("7. Tests et builds valides", level=1)
    add_code(doc, r'''cd "C:\Users\LOQ\Documents\yassmin pfe\mon_pfapp"
C:\Users\LOQ\dev\flutter\bin\cache\dart-sdk\bin\dart.exe analyze
C:\Users\LOQ\dev\flutter\bin\flutter.bat test
C:\Users\LOQ\dev\flutter\bin\flutter.bat build apk --debug --dart-define=DEMO_MODE=false --dart-define=API_BASE_URL=http://172.17.182.162:8000/api''')
    add_code(doc, r'''cd "C:\Users\LOQ\Documents\yassmin pfe\mon_pfapi"
..\tools\runtime\php-8.4.22\php.exe artisan test''')

    doc.add_heading("8. Guide de debug rapide", level=1)
    add_matrix(doc, ["Symptome", "Cause probable", "Solution"], [
        ["API ne repond pas", "Laravel serve non lance", "Relancer artisan serve port 8000."],
        ["Login impossible sur telephone", "Mauvaise IP PC ou Wi-Fi different", "Utiliser ipconfig et reconstruire/run avec la bonne IP."],
        ["ADB vide", "USB debug non accepte", "Accepter popup RSA, changer cable ou mode transfert fichier."],
        ["dart analyze bloque", "Wrapper dart.bat", "Utiliser dart.exe direct."],
        ["Windows atlstr.h", "Composant ATL absent", "Installer Visual Studio C++ ATL."],
        ["Android cleartext", "HTTP bloque en release", "Utiliser debug pour HTTP local ou HTTPS en release."],
    ], [1.75, 2.35, 2.4])

    doc.add_heading("9. Architecture et fichiers a comprendre", level=1)
    add_image(doc, diagrams["architecture"], "Architecture Flutter + Laravel.", width=6.3)
    add_matrix(doc, ["Fichier", "A retenir"], [
        ["lib/data/api_client.dart", "Centralise HTTP, token, JSON, timeout et erreurs."],
        ["lib/features/client/data/order_service.dart", "Cree et recupere les commandes."],
        ["mon_pfapi/routes/api.php", "Liste des endpoints REST."],
        ["mon_pfapi/app/Http/Middleware/TokenAuth.php", "Verifie token et roles."],
        ["mon_pfapi/database/seeders/DatabaseSeeder.php", "Jeu de donnees de demonstration."],
    ], [2.6, 3.9])

    doc.add_heading("10. Scenario de soutenance", level=1)
    add_numbered(doc, [
        "Presenter la problematique du restaurant.",
        "Lancer l'API Laravel et montrer /api/health.",
        "Ouvrir l'app en client avec yassmine@monpf.fr.",
        "Ajouter des plats au panier et creer une commande.",
        "Montrer le suivi de commande.",
        "Se connecter admin pour montrer stats et commandes.",
        "Se connecter livreur pour montrer la file de livraison.",
        "Terminer par les tests et les perspectives production.",
    ])
    add_image(doc, SCREENSHOT_DIR / "contact-sheet.png", "Captures principales pour la demonstration.", width=6.2)
    doc.save(TUTORIAL_DOCX)


def pdf_styles() -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle("title", parent=base["Title"], fontName="Helvetica-Bold", fontSize=24, leading=30, alignment=TA_CENTER, textColor=colors.HexColor("#1F2933"), spaceAfter=12),
        "subtitle": ParagraphStyle("subtitle", parent=base["Normal"], fontName="Helvetica-Bold", fontSize=12, leading=16, alignment=TA_CENTER, textColor=colors.HexColor("#1F4D78"), spaceAfter=14),
        "h1": ParagraphStyle("h1", parent=base["Heading1"], fontName="Helvetica-Bold", fontSize=15, leading=19, textColor=colors.HexColor("#2E74B5"), spaceBefore=12, spaceAfter=7),
        "h2": ParagraphStyle("h2", parent=base["Heading2"], fontName="Helvetica-Bold", fontSize=12, leading=15, textColor=colors.HexColor("#1F4D78"), spaceBefore=8, spaceAfter=5),
        "body": ParagraphStyle("body", parent=base["BodyText"], fontName="Helvetica", fontSize=9.5, leading=12.7, textColor=colors.HexColor("#1F2933"), alignment=TA_JUSTIFY, spaceAfter=6),
        "small": ParagraphStyle("small", parent=base["BodyText"], fontName="Helvetica", fontSize=8, leading=10, textColor=colors.HexColor("#4B5563"), spaceAfter=3),
        "caption": ParagraphStyle("caption", parent=base["Italic"], fontName="Helvetica-Oblique", fontSize=8, leading=10, alignment=TA_CENTER, textColor=colors.HexColor("#6B7280"), spaceAfter=6),
        "code": ParagraphStyle("code", parent=base["Code"], fontName="Courier", fontSize=7.5, leading=9.2, textColor=colors.white, backColor=colors.HexColor("#111827"), leftIndent=6, rightIndent=6, spaceBefore=4, spaceAfter=8),
    }


def pdf_para(text: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(sanitize(text), style)


def pdf_table(headers: list[str], rows: list[list[str]], widths_cm: list[float], styles: dict[str, ParagraphStyle]) -> Table:
    data = [[pdf_para(h, styles["small"]) for h in headers]]
    for row in rows:
        data.append([pdf_para(cell, styles["small"]) for cell in row])
    table = Table(data, colWidths=[w * cm for w in widths_cm], hAlign="CENTER")
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#1F4D78")),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#CBD5E1")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    return table


def pdf_image(path: Path, caption: str, styles: dict[str, ParagraphStyle], width_cm: float = 15.5) -> list:
    if not path.exists():
        return [pdf_para(f"[Image manquante: {path.name}]", styles["body"])]
    im = Image.open(path)
    ratio = im.height / im.width
    width = width_cm * cm
    height = width * ratio
    max_height = 13.5 * cm
    if height > max_height:
        height = max_height
        width = height / ratio
    return [RLImage(str(path), width=width, height=height), pdf_para(caption, styles["caption"])]


def pdf_bullets(items: list[str], styles: dict[str, ParagraphStyle]) -> ListFlowable:
    return ListFlowable([ListItem(pdf_para(item, styles["body"])) for item in items], bulletType="bullet", leftIndent=16)


def pdf_numbered(items: list[str], styles: dict[str, ParagraphStyle]) -> ListFlowable:
    return ListFlowable([ListItem(pdf_para(item, styles["body"])) for item in items], bulletType="1", leftIndent=18)


def footer(canvas, doc) -> None:
    canvas.saveState()
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(colors.HexColor("#6B7280"))
    canvas.drawString(1.8 * cm, 1.0 * cm, f"{INFO.project} - {INFO.student}")
    canvas.drawRightString(19.2 * cm, 1.0 * cm, f"Page {doc.page}")
    canvas.restoreState()


def build_report_pdf(diagrams: dict[str, Path]) -> None:
    styles = pdf_styles()
    story: list = []
    if LOGO_PATH.exists():
        story.extend(pdf_image(LOGO_PATH, "", styles, width_cm=4.3)[:1])
    story.extend([
        Spacer(1, 0.5 * cm),
        pdf_para(INFO.ministry, styles["subtitle"]),
        pdf_para(INFO.academy, styles["subtitle"]),
        pdf_para("Rapport de Projet de Fin d'Etudes", styles["title"]),
        pdf_para(INFO.formation, styles["subtitle"]),
        Spacer(1, 0.4 * cm),
        pdf_para(INFO.report_title, styles["title"]),
        Spacer(1, 0.5 * cm),
        pdf_table(["Champ", "Valeur"], [["Projet", INFO.project], ["Realise par", INFO.student], ["Etablissement", INFO.school], ["Annee", INFO.academic_year]], [4.5, 10.5], styles),
        PageBreak(),
    ])
    for heading, paragraphs in REPORT_SECTIONS:
        story.append(pdf_para(heading, styles["h1"]))
        for paragraph in paragraphs:
            story.append(pdf_para(paragraph, styles["body"]))
        story.append(PageBreak())
    story.extend([
        pdf_para("Table des matieres", styles["h1"]),
        pdf_numbered(["Introduction generale", "Contexte et problematique", "Cahier des charges", "Analyse et conception", "Architecture technique", "Realisation backend", "Realisation Flutter", "Securite et qualite", "Tests et validation", "Installation et mise en service", "Difficultes", "Conclusion et perspectives", "Annexes"], styles),
        PageBreak(),
    ])
    chapters = [
        ("1. Introduction generale", ["Mon PF App est un projet de fin d'etudes BTS DAI portant sur la digitalisation des commandes d'un restaurant. Le projet combine application mobile, backend API, base de donnees et documentation technique.", "La version livree vise une demonstration presque finale: elle est utilisable localement, testable, documentee et structuree professionnellement."]),
        ("2. Contexte et problematique", ["La gestion manuelle des commandes provoque erreurs, retards et manque de visibilite. Une application mobile permet de centraliser les informations et de mieux coordonner client, restaurant et livreur."]),
        ("3. Cahier des charges", ["Le cahier de charge demande les fonctions client, restaurant/admin, serveur et livreur, ainsi que des exigences de securite, rapidite, maintenance et ergonomie."]),
        ("4. Analyse et conception", ["Les principales entites sont User, MenuItem, Category, Order et OrderItem. Les roles determinent les actions autorisees."]),
        ("5. Architecture technique", ["Flutter gere l'interface et Laravel expose l'API REST. SQLite est utilisee pour la demonstration locale afin de faciliter l'installation." ]),
        ("6. Realisation backend", ["Le backend Laravel contient les migrations, seeders, controleurs API, middleware de role et tests PHPUnit."]),
        ("7. Realisation Flutter", ["Flutter contient les ecrans client, livreur et admin, les services API et les widgets partages. Les images de plats et transitions ameliorent l'experience utilisateur."]),
        ("8. Securite et qualite", ["Les roles sont imposes cote serveur, le token est stocke en secure storage cote app, les appels ont un timeout et les tests couvrent les flux principaux."]),
        ("9. Tests et validation", ["Les tests Flutter et Laravel passent. Un test live API a cree une commande reelle via POST /api/orders."]),
        ("10. Installation et mise en service", ["Le backend se lance avec artisan serve. Flutter se lance avec DEMO_MODE=false et API_BASE_URL. Android debug peut utiliser HTTP local; release doit viser HTTPS."]),
        ("11. Difficultes rencontrees", ["Les principaux blocages ont ete les permissions Chocolatey, le wrapper dart.bat, la politique cleartext Android et le composant ATL Windows manquant."]),
        ("12. Conclusion et perspectives", ["Le projet valide le parcours demo-prod. Les evolutions concernent hebergement HTTPS, base hebergee, Sanctum, notifications, paiement et CI/CD."]),
    ]
    for title, paras in chapters:
        story.append(pdf_para(title, styles["h1"]))
        for para in paras:
            story.append(pdf_para(para, styles["body"]))
        if "Architecture" in title:
            story.extend(pdf_image(diagrams["architecture"], "Architecture demo-prod.", styles))
        if "Analyse" in title:
            story.extend(pdf_image(diagrams["data_model"], "Modele logique simplifie.", styles))
        if "Tests" in title:
            story.extend(pdf_image(diagrams["order_cycle"], "Cycle de commande.", styles))
        if "Realisation Flutter" in title and (SCREENSHOT_DIR / "contact-sheet.png").exists():
            story.extend(pdf_image(SCREENSHOT_DIR / "contact-sheet.png", "Captures de l'application.", styles, width_cm=14.2))
        story.append(Spacer(1, 0.25 * cm))
    story.append(pdf_para("Annexes - Comptes de test", styles["h1"]))
    story.append(pdf_table(["Email", "Role", "Mot de passe"], [["yassmine@monpf.fr", "client", "password"], ["admin@monpf.fr", "admin", "password"], ["serveur@monpf.fr", "serveur", "password"], ["livreur@monpf.fr", "livreur", "password"]], [6.0, 4.0, 4.0], styles))
    story.append(pdf_para("Commandes principales", styles["h2"]))
    story.append(Paragraph(sanitize(r'''cd "C:\Users\LOQ\Documents\yassmin pfe\mon_pfapi"
..\tools\runtime\php-8.4.22\php.exe artisan serve --host=0.0.0.0 --port=8000

cd "C:\Users\LOQ\Documents\yassmin pfe\mon_pfapp"
C:\Users\LOQ\dev\flutter\bin\flutter.bat test
C:\Users\LOQ\dev\flutter\bin\cache\dart-sdk\bin\dart.exe analyze''').replace("\n", "<br/>"), styles["code"]))
    doc = SimpleDocTemplate(str(PFE_PDF), pagesize=A4, rightMargin=1.7 * cm, leftMargin=1.7 * cm, topMargin=1.65 * cm, bottomMargin=1.55 * cm, title="Rapport PFE - Mon PF App")
    doc.build(story, onFirstPage=footer, onLaterPages=footer)


def build_tutorial_pdf(diagrams: dict[str, Path]) -> None:
    styles = pdf_styles()
    story: list = []
    if LOGO_PATH.exists():
        story.extend(pdf_image(LOGO_PATH, "", styles, width_cm=3.5)[:1])
    story.extend([
        pdf_para("Mon PF App", styles["title"]),
        pdf_para("Tutoriel complet d'installation, test, debug, build et demonstration", styles["subtitle"]),
        pdf_para(f"{INFO.student} - {INFO.formation}", styles["subtitle"]),
        PageBreak(),
        pdf_para("1. Ce que contient le projet", styles["h1"]),
        pdf_para("Le repo contient mon_pfapp pour Flutter, mon_pfapi pour Laravel, docs pour la documentation et tools pour les scripts.", styles["body"]),
        pdf_table(["Dossier", "Role"], [["mon_pfapp", "Application Flutter"], ["mon_pfapi", "Backend Laravel"], ["docs", "Rapports, runbooks et specifications"], ["tools", "Generation documents et runtime local"]], [4.0, 11.0], styles),
        pdf_para("2. Lancer le backend", styles["h1"]),
        Paragraph(sanitize(r'''cd "C:\Users\LOQ\Documents\yassmin pfe\mon_pfapi"
..\tools\runtime\php-8.4.22\php.exe artisan migrate:fresh --seed
..\tools\runtime\php-8.4.22\php.exe artisan serve --host=0.0.0.0 --port=8000''').replace("\n", "<br/>"), styles["code"]),
        pdf_para("3. Lancer Flutter", styles["h1"]),
        Paragraph(sanitize(r'''cd "C:\Users\LOQ\Documents\yassmin pfe\mon_pfapp"
C:\Users\LOQ\dev\flutter\bin\flutter.bat run --dart-define=DEMO_MODE=false --dart-define=API_BASE_URL=http://127.0.0.1:8000/api''').replace("\n", "<br/>"), styles["code"]),
        pdf_para("4. Android", styles["h1"]),
        pdf_numbered(["Activer Options developpeur.", "Activer Debogage USB.", "Brancher le telephone.", "Accepter la popup RSA.", "Installer app-debug.apk ou lancer flutter run."], styles),
        pdf_para("5. Tests", styles["h1"]),
        Paragraph(sanitize(r'''C:\Users\LOQ\dev\flutter\bin\cache\dart-sdk\bin\dart.exe analyze
C:\Users\LOQ\dev\flutter\bin\flutter.bat test
..\tools\runtime\php-8.4.22\php.exe artisan test''').replace("\n", "<br/>"), styles["code"]),
        pdf_para("6. Debug rapide", styles["h1"]),
        pdf_table(["Symptome", "Solution"], [["API ne repond pas", "Relancer artisan serve."], ["Telephone ne se connecte pas", "Verifier IP Wi-Fi PC et backend."], ["ADB vide", "Accepter RSA / changer cable."], ["Windows atlstr.h", "Installer Visual Studio C++ ATL."], ["dart bloque", "Utiliser dart.exe direct."]], [5.0, 10.0], styles),
        pdf_para("7. Architecture", styles["h1"]),
    ])
    story.extend(pdf_image(diagrams["architecture"], "Architecture Flutter + Laravel.", styles))
    story.append(pdf_para("8. Scenario de soutenance", styles["h1"]))
    story.append(pdf_numbered(["Presenter la problematique.", "Lancer l'API.", "Se connecter client.", "Creer une commande.", "Montrer admin et livreur.", "Montrer les tests et limites production."], styles))
    if (SCREENSHOT_DIR / "contact-sheet.png").exists():
        story.extend(pdf_image(SCREENSHOT_DIR / "contact-sheet.png", "Captures principales.", styles, width_cm=14.2))
    doc = SimpleDocTemplate(str(TUTORIAL_PDF), pagesize=A4, rightMargin=1.7 * cm, leftMargin=1.7 * cm, topMargin=1.65 * cm, bottomMargin=1.55 * cm, title="Tutoriel Mon PF App")
    doc.build(story, onFirstPage=footer, onLaterPages=footer)


def build_report_pdf_full(diagrams: dict[str, Path]) -> None:
    styles = pdf_styles()
    story: list = []

    def page_title(title: str, subtitle: str | None = None) -> None:
        story.append(pdf_para(title, styles["h1"]))
        if subtitle:
            story.append(pdf_para(subtitle, styles["body"]))
        story.append(Spacer(1, 0.15 * cm))

    def chapter(title: str, paragraphs: list[str], *, image: Path | None = None, table: tuple[list[str], list[list[str]], list[float]] | None = None, bullets: list[str] | None = None, code: str | None = None) -> None:
        page_title(title)
        for paragraph in paragraphs:
            story.append(pdf_para(paragraph, styles["body"]))
        if bullets:
            story.append(pdf_bullets(bullets, styles))
        if table:
            headers, rows, widths = table
            story.append(pdf_table(headers, rows, widths, styles))
        if code:
            story.append(Paragraph(sanitize(code).replace("\n", "<br/>"), styles["code"]))
        if image:
            story.extend(pdf_image(image, f"Illustration - {title}", styles))
        story.append(PageBreak())

    if LOGO_PATH.exists():
        story.extend(pdf_image(LOGO_PATH, "", styles, width_cm=4.4)[:1])
    story.extend([
        Spacer(1, 0.25 * cm),
        pdf_para(INFO.ministry, styles["subtitle"]),
        pdf_para(INFO.academy, styles["subtitle"]),
        pdf_para(INFO.direction, styles["subtitle"]),
        Spacer(1, 0.25 * cm),
        pdf_para("Rapport de Projet de Fin d'Etudes", styles["title"]),
        pdf_para("Deuxieme annee BTS - Developpement des Applications Informatiques", styles["subtitle"]),
        Spacer(1, 0.45 * cm),
        pdf_para(INFO.report_title, styles["title"]),
        Spacer(1, 0.35 * cm),
        pdf_table(["Champ", "Valeur"], [
            ["Projet", INFO.project],
            ["Realise par", INFO.student],
            ["Etablissement", INFO.school],
            ["Encadrement", INFO.supervisor],
            ["Annee de formation", INFO.academic_year],
        ], [4.3, 10.7], styles),
        PageBreak(),
    ])

    chapter("Remerciements", [
        "Je remercie l'equipe pedagogique du BTS Developpement des Applications Informatiques du Lycee Qualifiant Technique Ibn Al-Haitam pour l'encadrement, les connaissances techniques et la methodologie acquises durant la formation.",
        "Je remercie egalement toutes les personnes qui ont participe a la definition du besoin, aux tests, aux retours d'utilisation et a l'amelioration progressive de l'application Mon PF App.",
        "Ce projet a permis de rapprocher les acquis de programmation, base de donnees, analyse UML, securite et documentation d'un besoin applicatif concret: la gestion des commandes d'un restaurant.",
    ])

    chapter("Resume", [
        "Ce rapport presente Mon PF App, une application mobile de gestion des commandes et des livraisons d'un restaurant. Le projet repond a la problematique de la gestion manuelle: erreurs de commande, retards de preparation, manque de suivi et communication insuffisante entre client, restaurant et livreur.",
        "La solution s'appuie sur Flutter pour l'application mobile et Laravel pour l'API backend. Le systeme permet au client de consulter le menu, ajouter des plats au panier, passer une commande et suivre son statut. Il propose aussi des espaces dedies a l'administrateur, au serveur et au livreur.",
        "La version livree est une version demo-prod: elle dispose d'une interface aboutie, d'un backend local fonctionnel, d'une base SQLite seedee, de tests automatiques et d'une documentation permettant d'installer, tester, debugger et presenter le projet.",
    ])

    story.append(pdf_para("Table des matieres", styles["h1"]))
    story.append(pdf_numbered([
        "Introduction generale",
        "Contexte et problematique",
        "Objectifs du projet",
        "Cahier des charges",
        "Acteurs du systeme",
        "Analyse fonctionnelle",
        "Conception des donnees",
        "Architecture technique",
        "Backend Laravel",
        "Application Flutter",
        "Securite",
        "Algorithmes et regles metier",
        "Interface et experience utilisateur",
        "Tests et validation",
        "Mise en service locale",
        "Android et Windows",
        "Difficultes et solutions",
        "Validation du cahier de charge",
        "Perspectives",
        "Conclusion",
        "Annexes",
    ], styles))
    story.append(PageBreak())

    chapter("1. Introduction generale", [
        "Dans le cadre de la formation BTS Developpement des Applications Informatiques, le projet de fin d'etudes doit demontrer la capacite a analyser un besoin, concevoir une solution, realiser une application, tester son fonctionnement et documenter le travail effectue.",
        "Mon PF App est une application de gestion des commandes d'un restaurant. Elle montre un parcours client complet et des espaces operationnels pour l'administration et la livraison.",
        "Le choix du sujet est pertinent pour un PFE car il mobilise plusieurs competences: analyse des acteurs, modelisation des donnees, interface mobile, API REST, securite, tests, build Android et documentation professionnelle.",
    ], table=(["Element", "Valeur"], [["Formation", INFO.formation], ["Projet", INFO.project], ["Cible principale", "Android et Windows"], ["Type", "Application mobile + backend API"]], [4.0, 11.0]))

    chapter("2. Contexte et problematique", [
        "Les restaurants gerent de plus en plus de commandes a emporter et en livraison. Lorsque les commandes sont prises manuellement, les erreurs de saisie et les oublis deviennent frequents.",
        "Le client souhaite connaitre l'etat de sa commande. Le restaurant doit organiser les preparations. Le livreur doit savoir quelles commandes sont disponibles et lesquelles lui sont affectees.",
        "La problematique du projet est donc: comment digitaliser le processus de commande afin de reduire les erreurs humaines, accelerer le traitement et ameliorer la communication entre les acteurs ?",
    ], table=(["Probleme", "Consequence", "Reponse Mon PF App"], [["Erreur de commande", "Plat incorrect ou quantite fausse", "Panier numerique et validation"], ["Retard", "Client insatisfait", "Suivi par statut"], ["Mauvaise communication", "Livreur ou restaurant mal informe", "Espaces dedies"], ["Menu difficile a maintenir", "Prix/disponibilite non a jour", "Catalogue API et CRUD admin"]], [4.0, 5.0, 6.0]))

    chapter("3. Objectifs du projet", [
        "L'objectif general est de realiser une application mobile de gestion des commandes d'un restaurant. Elle doit organiser le processus depuis le choix des plats jusqu'a la livraison.",
        "Les objectifs specifiques consistent a permettre au client de consulter le menu, de passer une commande et de suivre son etat; a permettre au restaurant de gerer les commandes; et a permettre au livreur de consulter et accepter les livraisons.",
        "Sur le plan technique, le projet doit etre structure, maintenable et testable. Le repository a donc ete organise avec un frontend Flutter, un backend Laravel, des documents de specification, des guides et des scripts.",
    ], bullets=["Digitaliser la commande client.", "Centraliser le menu et les commandes.", "Gerer les roles client, admin, serveur et livreur.", "Construire un APK Android testable.", "Documenter installation, debug et soutenance."])

    chapter("4. Cahier des charges", [
        "Le cahier de charge fourni indique les besoins fonctionnels cote client, restaurant et livreur. Il indique aussi les besoins non fonctionnels: securite, simplicite, rapidite, fiabilite et maintenance.",
        "La version actuelle couvre le coeur du cahier de charge. Le client peut se connecter, consulter le menu, ajouter au panier, passer une commande et suivre son statut. Le livreur peut consulter la file et accepter une livraison. L'administrateur dispose de statistiques et d'une vision commandes.",
        "Certaines evolutions restent necessaires pour une vraie production commerciale: hebergement HTTPS, base hebergee, notifications, paiement et interface CRUD admin complete.",
    ], table=(["Besoin", "Etat"], [["Connexion/inscription", "Couvert demo + API"], ["Menu", "Couvert demo + API"], ["Panier/commande", "Couvert avec POST /api/orders"], ["Suivi", "Couvert visuellement; temps reel a ameliorer"], ["Livreur", "File et acceptation disponibles"], ["Admin", "Stats et endpoints disponibles"]], [6.5, 8.5]))

    chapter("5. Acteurs du systeme", [
        "Le systeme est concu autour de quatre acteurs. Chaque acteur a un role precis afin de respecter les responsabilites metier.",
        "Le client utilise l'application pour commander. L'administrateur gere le systeme global. Le serveur intervient dans le traitement des commandes cote restaurant. Le livreur prend en charge la livraison.",
    ], table=(["Acteur", "Responsabilites"], [["Client", "Consulter menu, panier, commander, suivre"], ["Administrateur", "Statistiques, commandes, menu, affectation"], ["Serveur", "Traitement operationnel des commandes"], ["Livreur", "Consulter, accepter, livrer"]], [4.2, 10.8]))

    chapter("6. Analyse fonctionnelle", [
        "L'analyse fonctionnelle decompose le systeme en cas d'utilisation. Chaque cas d'utilisation correspond a une action visible dans l'application ou dans l'API.",
        "Le parcours principal est le passage de commande. Le client choisit des plats, valide le panier, l'API cree une commande et les espaces operationnels peuvent ensuite suivre ou traiter cette commande.",
        "Cette analyse a guide la structure des fichiers Flutter et Laravel: les fonctionnalites sont regroupees par domaine, ce qui facilite la comprehension par Yassmine et par le jury.",
    ], image=diagrams["order_cycle"])

    chapter("7. Conception des donnees", [
        "Le modele de donnees s'organise autour des utilisateurs, categories, plats, commandes et lignes de commande. Cette structure evite de stocker une commande comme simple texte et permet de garder l'historique des prix et quantites.",
        "Le champ role de l'utilisateur permet de distinguer client, admin, serveur et livreur. La commande contient son statut, son adresse, son total et eventuellement le livreur affecte.",
        "SQLite est utilisee pour la demo locale. En production, le meme schema peut etre porte vers MySQL ou PostgreSQL.",
    ], image=diagrams["data_model"])

    chapter("8. Architecture technique", [
        "L'architecture separe le frontend du backend. Flutter gere l'interface, l'etat utilisateur et le panier. Laravel gere les validations, les routes REST et la persistance.",
        "Cette separation est professionnelle: elle permet de faire evoluer l'interface sans changer la base, et de remplacer SQLite par une base hebergee sans modifier toute l'application mobile.",
        "Le mode demo reste disponible pour garantir une presentation meme si le serveur n'est pas accessible. Le mode backend reel s'active avec les variables DEMO_MODE=false et API_BASE_URL.",
    ], image=diagrams["architecture"])

    chapter("9. Backend Laravel", [
        "Le backend se trouve dans mon_pfapi. Il expose des endpoints REST publics et proteges. Les endpoints publics sont /api/health, /api/menu, /api/register et /api/login.",
        "Les routes proteges utilisent un middleware TokenAuth. Le token est envoye dans le header Authorization: Bearer. Le backend stocke seulement un hash du token, ce qui reduit le risque en cas de fuite de base.",
        "Les tests PHPUnit valident le menu seed, la connexion, l'inscription forcee en client et la creation de commande.",
    ], table=(["Composant", "Role"], [["AuthController", "Register, login, logout"], ["MenuController", "Catalogue et CRUD menu"], ["OrderController", "Commandes, statuts, livreur, stats"], ["TokenAuth", "Controle token et role"], ["DatabaseSeeder", "Donnees demo"]], [5.0, 10.0]))

    chapter("10. Application Flutter", [
        "L'application Flutter se trouve dans mon_pfapp. Elle contient des ecrans pour l'authentification, l'accueil client, le menu, le panier, le suivi, le profil, l'administration et la livraison.",
        "La classe MonPfApp centralise l'etat principal: utilisateur connecte, ecran courant, menu charge, commandes et panier. Les services MenuService et OrderService recuperent les donnees depuis l'API ou depuis DemoData.",
        "L'interface utilise des images reelles de plats, un logo personnalise et des transitions animees afin d'obtenir un rendu plus proche d'un produit final.",
    ], image=SCREENSHOT_DIR / "contact-sheet.png")

    chapter("11. Securite", [
        "La securite a ete renforcee a plusieurs niveaux. La selection de role a ete retiree de l'inscription publique. Le role final est impose par le backend.",
        "Le token est stocke dans flutter_secure_storage cote application. Cote API, le token brut n'est pas conserve: seul son hash SHA-256 est stocke.",
        "L'APK release garde une politique reseau stricte. HTTP local est autorise seulement pour debug/profile afin de faciliter les tests sur Wi-Fi pendant la soutenance.",
    ], table=(["Risque", "Traitement"], [["Role public admin", "Role client force cote serveur"], ["Token local", "Secure storage"], ["Token base", "Hash du token"], ["HTTP release", "Refuse sauf debug/profile"], ["Serveur lent", "Timeout et erreurs"]], [6.0, 9.0]))

    chapter("12. Algorithmes et regles metier", [
        "Le projet repose sur des algorithmes applicatifs simples et defendables. Le panier verifie si un plat existe deja; si oui, il augmente la quantite, sinon il ajoute une nouvelle ligne.",
        "Le total est calcule par addition des lignes: prix unitaire multiplie par quantite, puis ajout des frais de livraison. La creation de commande refuse un panier vide.",
        "Le cycle de statut permet de representer l'avancement: preparation, pret, en route, livre ou annule. Les roles controlent les statuts qu'ils peuvent modifier.",
    ], table=(["Regle", "Explication"], [["Panier non vide", "Une commande exige au moins une ligne"], ["Quantite", "Borne entre 1 et 99 cote app"], ["Total", "Somme lignes + livraison"], ["Role", "Autorisation cote middleware"], ["Statut", "Normalisation cote backend"]], [5.0, 10.0]))

    chapter("13. Interface et experience utilisateur", [
        "L'interface a ete inspiree par la maquette restaurant fournie. Elle privilegie une presentation mobile claire: header rouge, cartes blanches, navigation basse, visuels de plats et boutons d'action visibles.",
        "L'objectif UI/UX est de permettre au client de comprendre rapidement le parcours. Le menu est filtre par categorie, le panier resume le total et le suivi affiche les etapes de livraison.",
        "Les espaces admin et livreur utilisent des cartes de statistiques et des listes de commandes afin de rester lisibles sur telephone.",
    ], image=SCREENSHOT_DIR / "02-home.png")

    chapter("14. Tests et validation", [
        "Les tests automatiques prouvent que l'application n'est pas seulement une maquette. Flutter teste les ecrans critiques et Laravel teste les flux API.",
        "Les commandes validees sont dart analyze, flutter test et php artisan test. Un test HTTP live a confirme qu'un client peut se connecter et creer une commande via l'API.",
        "Le build Android debug a ete genere et pointe vers l'API locale du PC. Le build Windows reste bloque par ATL, un composant externe Visual Studio.",
    ], table=(["Controle", "Resultat"], [["Dart analyze", "No issues found"], ["Flutter tests", "5 tests passed"], ["Laravel tests", "5 tests passed"], ["API live", "Commande creee"], ["APK debug", "Construit"], ["Windows", "ATL manquant"]], [5.5, 9.5]))

    chapter("15. Mise en service locale", [
        "La mise en service locale se fait en deux etapes: demarrer Laravel puis lancer Flutter avec l'URL API. Le backend doit rester ouvert pendant le test.",
        "Pour Android physique, le telephone doit etre sur le meme Wi-Fi que le PC, et l'URL doit utiliser l'adresse IPv4 Wi-Fi du PC au lieu de 127.0.0.1.",
        "Les commandes exactes sont documentees dans le tutoriel technique et le runbook demo-prod.",
    ], code=r'''cd "C:\Users\LOQ\Documents\yassmin pfe\mon_pfapi"
..\tools\runtime\php-8.4.22\php.exe artisan serve --host=0.0.0.0 --port=8000

cd "C:\Users\LOQ\Documents\yassmin pfe\mon_pfapp"
C:\Users\LOQ\dev\flutter\bin\flutter.bat run --dart-define=DEMO_MODE=false --dart-define=API_BASE_URL=http://127.0.0.1:8000/api''')

    chapter("16. Android et Windows", [
        "Android est la cible la plus avancee. L'APK debug existe et peut etre installe avec ADB lorsque le telephone est connecte et autorise.",
        "Windows est une cible prioritaire mais le build necessite l'installation du composant Visual Studio ATL. L'erreur exacte est atlstr.h introuvable.",
        "Cette distinction doit etre expliquee clairement au jury: le code Windows n'est pas bloque par Flutter, mais par une dependance systeme manquante.",
    ], table=(["Cible", "Etat", "Action"], [["Android", "APK debug pret", "Installer via ADB"], ["Windows", "ATL manquant", "Installer composant Visual Studio"], ["Release Android", "HTTPS requis", "Deployer API HTTPS"]], [3.5, 5.0, 6.5]))

    chapter("17. Difficultes et solutions", [
        "Plusieurs difficultes techniques ont ete rencontrees pendant la transformation demo en demo-prod. Chaque difficulte a ete diagnostiquee puis contournee proprement.",
        "PHP et Composer n'etaient pas installes globalement. Chocolatey etait bloque par des permissions Windows. La solution a ete une installation portable locale dans tools/runtime.",
        "Le wrapper dart.bat etait lent dans ce contexte. L'analyse est donc lancee avec dart.exe direct. Le build Android a aussi necessite une correction du manifest debug pour autoriser HTTP local.",
    ], table=(["Difficulte", "Solution"], [["Chocolatey bloque", "PHP/Composer portable"], ["dart.bat lent", "dart.exe direct"], ["Cleartext Android", "tools:replace debug/profile"], ["ADB non trouve", "Chemin complet platform-tools"], ["Windows ATL", "Installer composant Visual Studio"]], [6.0, 9.0]))

    chapter("18. Validation du cahier de charge", [
        "La version actuelle valide les besoins essentiels du cahier de charge pour une demonstration PFE. Les parcours principaux sont presents et connectes a une API locale.",
        "La validation complete production demanderait toutefois un serveur HTTPS heberge, une base MySQL/PostgreSQL et une synchronisation temps reel. Ces limites sont normales pour une version demo-prod.",
    ], table=(["Exigence", "Validation"], [["Client menu/panier/commande", "Oui"], ["Suivi commande", "Oui, visuel + statut"], ["Admin statistiques", "Oui"], ["Livreur file/acceptation", "Oui"], ["CRUD menu", "API oui, UI a completer"], ["Temps reel", "Evolution"], ["HTTPS deploye", "Evolution"]], [6.5, 8.5]))

    chapter("19. Perspectives", [
        "Les perspectives principales concernent la production reelle. L'API devrait etre deployee sur un domaine HTTPS et connectee a une base hebergee.",
        "Laravel Sanctum peut remplacer le token simple pour une gestion d'authentification plus standard. Des notifications push peuvent informer le client quand le statut change.",
        "Une interface admin complete pour le CRUD menu permettrait au gerant de modifier plats et disponibilites sans intervention technique.",
    ], bullets=["Deploiement HTTPS", "Base MySQL/PostgreSQL", "Laravel Sanctum", "Notifications push", "Paiement en ligne", "CI/CD GitHub Actions"])

    chapter("20. Conclusion", [
        "Mon PF App constitue une base solide de projet de fin d'etudes BTS. Le projet montre la capacite a comprendre un besoin, concevoir une architecture, developper une application, ajouter un backend, securiser les roles et tester les flux.",
        "La version demo-prod est suffisante pour une presentation claire: elle possede un parcours client, un espace admin, un espace livreur, une API Laravel, une base SQLite, un APK Android et une documentation detaillee.",
        "Le travail restant est bien identifie et peut etre presente comme perspectives de production. Cette transparence renforce la credibilite du projet.",
    ])

    story.append(pdf_para("Annexe A - Comptes de demonstration", styles["h1"]))
    story.append(pdf_table(["Email", "Role", "Mot de passe"], [["yassmine@monpf.fr", "client", "password"], ["admin@monpf.fr", "admin", "password"], ["serveur@monpf.fr", "serveur", "password"], ["livreur@monpf.fr", "livreur", "password"]], [6.2, 4.0, 4.2], styles))
    story.append(PageBreak())
    story.append(pdf_para("Annexe B - Endpoints API", styles["h1"]))
    story.append(pdf_table(["Methode", "Endpoint", "Role"], [["GET", "/api/health", "public"], ["GET", "/api/menu", "public"], ["POST", "/api/register", "public"], ["POST", "/api/login", "public"], ["GET", "/api/orders", "connecte"], ["POST", "/api/orders", "client"], ["GET", "/api/admin/stats", "admin/serveur"], ["GET", "/api/driver/orders", "livreur/admin"]], [3.0, 7.0, 4.5], styles))
    story.append(PageBreak())
    story.append(pdf_para("Annexe C - Fichiers importants", styles["h1"]))
    story.append(pdf_table(["Fichier", "Role"], [["mon_pfapp/lib/app/mon_pf_app.dart", "Etat global et navigation"], ["mon_pfapp/lib/data/api_client.dart", "Client HTTP"], ["mon_pfapi/routes/api.php", "Routes REST"], ["mon_pfapi/database/seeders/DatabaseSeeder.php", "Donnees demo"], ["docs/demo-prod-runbook.md", "Guide lancement"], ["docs/backend-api.md", "Documentation API"]], [7.0, 8.0], styles))

    doc = SimpleDocTemplate(str(PFE_PDF), pagesize=A4, rightMargin=1.7 * cm, leftMargin=1.7 * cm, topMargin=1.65 * cm, bottomMargin=1.55 * cm, title="Rapport PFE - Mon PF App")
    doc.build(story, onFirstPage=footer, onLaterPages=footer)


def build_tutorial_pdf_full(diagrams: dict[str, Path]) -> None:
    styles = pdf_styles()
    story: list = []

    def section(title: str, paragraphs: list[str], *, table: tuple[list[str], list[list[str]], list[float]] | None = None, code: str | None = None, image: Path | None = None, bullets: list[str] | None = None, numbered: list[str] | None = None) -> None:
        story.append(pdf_para(title, styles["h1"]))
        for paragraph in paragraphs:
            story.append(pdf_para(paragraph, styles["body"]))
        if bullets:
            story.append(pdf_bullets(bullets, styles))
        if numbered:
            story.append(pdf_numbered(numbered, styles))
        if table:
            headers, rows, widths = table
            story.append(pdf_table(headers, rows, widths, styles))
        if code:
            story.append(Paragraph(sanitize(code).replace("\n", "<br/>"), styles["code"]))
        if image:
            story.extend(pdf_image(image, f"Illustration - {title}", styles))
        story.append(PageBreak())

    if LOGO_PATH.exists():
        story.extend(pdf_image(LOGO_PATH, "", styles, width_cm=3.5)[:1])
    story.extend([
        pdf_para("Mon PF App", styles["title"]),
        pdf_para("Tutoriel technique complet", styles["subtitle"]),
        pdf_para("Installer, lancer, tester, debugger, builder et presenter l'application", styles["subtitle"]),
        pdf_para(f"{INFO.student} - {INFO.formation}", styles["subtitle"]),
        PageBreak(),
        pdf_para("Sommaire", styles["h1"]),
        pdf_numbered([
            "Vue d'ensemble du projet",
            "Prerequis et chemins importants",
            "Backend Laravel",
            "Flutter en mode demo",
            "Flutter avec backend reel",
            "Android telephone",
            "Windows",
            "Tests automatiques",
            "Builds",
            "Debug API et mobile",
            "Architecture et fichiers importants",
            "Comptes de test",
            "Scenario de soutenance",
            "Checklist finale",
        ], styles),
        PageBreak(),
    ])

    section("1. Vue d'ensemble du projet", [
        "Mon PF App contient une application Flutter et un backend Laravel. Flutter gere l'interface mobile; Laravel gere l'API REST et la base de demonstration.",
        "Le projet peut fonctionner en mode demo sans serveur ou en mode backend reel avec DEMO_MODE=false et API_BASE_URL. Pour la soutenance, garder les deux modes est important: le mode demo securise la presentation, le mode API prouve le fonctionnement reel.",
    ], table=(["Dossier", "Role"], [["mon_pfapp", "Frontend Flutter"], ["mon_pfapi", "Backend Laravel API"], ["docs", "Rapports, tutoriels, cahier de charge"], ["tools", "Scripts et runtime local portable"]], [4.0, 11.0]))

    section("2. Prerequis et chemins importants", [
        "Ces chemins sont ceux valides sur cette machine. Ils evitent les problemes de PATH et les commandes qui prennent trop de temps.",
        "Point important: pour l'analyse Dart, utiliser dart.exe direct. Sur cette machine, le wrapper dart.bat peut bloquer ou prendre beaucoup trop de temps.",
    ], table=(["Outil", "Chemin"], [["Flutter", r"C:\Users\LOQ\dev\flutter\bin\flutter.bat"], ["Dart", r"C:\Users\LOQ\dev\flutter\bin\cache\dart-sdk\bin\dart.exe"], ["PHP", r"C:\Users\LOQ\Documents\yassmin pfe\tools\runtime\php-8.4.22\php.exe"], ["ADB", r"C:\Users\LOQ\AppData\Local\Android\Sdk\platform-tools\adb.exe"], ["Repo", r"C:\Users\LOQ\Documents\yassmin pfe"]], [3.3, 11.7]))

    section("3. Lancer le backend Laravel", [
        "Le backend doit etre lance avant Flutter si l'application est en mode backend reel. La commande migrate:fresh --seed remet la base de demonstration a zero et recree les comptes de test.",
        "Laisser la fenetre artisan serve ouverte pendant les tests.",
    ], code=r'''cd "C:\Users\LOQ\Documents\yassmin pfe\mon_pfapi"
..\tools\runtime\php-8.4.22\php.exe artisan migrate:fresh --seed
..\tools\runtime\php-8.4.22\php.exe artisan serve --host=0.0.0.0 --port=8000''')

    section("4. Verifier l'API", [
        "Avant d'ouvrir Flutter, verifier que l'API repond. /api/health doit retourner status=ok. /api/menu doit retourner les categories et les plats.",
        "Si ces commandes echouent, le probleme vient du backend ou du port 8000, pas de Flutter.",
    ], code=r'''Invoke-RestMethod http://127.0.0.1:8000/api/health
Invoke-RestMethod http://127.0.0.1:8000/api/menu''')

    section("5. Lancer Flutter en mode demo", [
        "Le mode demo ne demande aucun serveur. Il est utile pour montrer l'interface meme si l'API n'est pas lancee.",
        "Ce mode utilise DemoData dans l'application. Il permet de tester login, accueil, menu, panier, suivi, profil, admin et livreur.",
    ], code=r'''cd "C:\Users\LOQ\Documents\yassmin pfe\mon_pfapp"
C:\Users\LOQ\dev\flutter\bin\flutter.bat run''')

    section("6. Lancer Flutter avec backend reel", [
        "En mode backend reel, l'application utilise l'API Laravel. Sur Windows local, utiliser 127.0.0.1. Sur telephone, utiliser l'adresse IP Wi-Fi du PC.",
        "La variable DEMO_MODE=false indique a Flutter de ne pas utiliser les donnees demo.",
    ], code=r'''cd "C:\Users\LOQ\Documents\yassmin pfe\mon_pfapp"
C:\Users\LOQ\dev\flutter\bin\flutter.bat run -d windows --dart-define=DEMO_MODE=false --dart-define=API_BASE_URL=http://127.0.0.1:8000/api''')

    section("7. Tester sur telephone Android", [
        "Pour tester sur un vrai telephone, il faut que le backend soit lance sur le PC et que le telephone puisse joindre l'adresse IP du PC.",
        "Le build debug Android autorise HTTP local. Le build release doit utiliser HTTPS.",
    ], numbered=["Activer Options developpeur.", "Activer Debogage USB.", "Brancher le telephone en USB.", "Accepter la popup RSA.", "Verifier avec adb devices.", "Installer app-debug.apk ou lancer flutter run."], code=r'''C:\Users\LOQ\AppData\Local\Android\Sdk\platform-tools\adb.exe devices
C:\Users\LOQ\AppData\Local\Android\Sdk\platform-tools\adb.exe install -r build\app\outputs\flutter-apk\app-debug.apk''')

    section("8. APK debug actuel", [
        "Un APK debug a ete construit pour l'API locale de cette machine. Si l'adresse IP Wi-Fi change, il faut reconstruire l'APK avec la nouvelle URL API.",
    ], table=(["Element", "Valeur"], [["APK", r"C:\Users\LOQ\Documents\yassmin pfe\mon_pfapp\build\app\outputs\flutter-apk\app-debug.apk"], ["API cible actuelle", "http://172.17.182.162:8000/api"], ["Taille approximative", "174 MB debug"], ["Usage", "Test local sur telephone, pas publication"]], [4.0, 11.0]))

    section("9. Build Android", [
        "Le build debug est recommande pour tester l'API locale HTTP. Le build release est recommande pour une version finale, mais il doit pointer vers une API HTTPS.",
    ], code=r'''C:\Users\LOQ\dev\flutter\bin\flutter.bat build apk --debug --dart-define=DEMO_MODE=false --dart-define=API_BASE_URL=http://172.17.182.162:8000/api

C:\Users\LOQ\dev\flutter\bin\flutter.bat build apk --release --dart-define=DEMO_MODE=false --dart-define=API_BASE_URL=https://votre-domaine/api''')

    section("10. Windows", [
        "Windows est une cible importante pour la demonstration sur PC, mais le build est actuellement bloque par le composant ATL manquant dans Visual Studio.",
        "Erreur observee: fatal error C1083: impossible d'ouvrir le fichier include atlstr.h.",
    ], numbered=["Ouvrir Visual Studio Installer.", "Modifier Visual Studio Community 2026.", "Chercher ATL dans Individual components.", "Installer C++ ATL for latest v143/v14x build tools ou equivalent.", "Relancer flutter build windows."], code=r'''C:\Users\LOQ\dev\flutter\bin\flutter.bat build windows --debug --dart-define=DEMO_MODE=false --dart-define=API_BASE_URL=http://127.0.0.1:8000/api''')

    section("11. Tests automatiques", [
        "Les tests prouvent que les parcours principaux fonctionnent. Ils sont utiles pour la soutenance car ils montrent une demarche professionnelle.",
        "Les tests Flutter couvrent login, inscription sans role public, menu, panier, suivi, admin et livreur. Les tests Laravel couvrent menu, login, role client force et creation commande.",
    ], code=r'''cd "C:\Users\LOQ\Documents\yassmin pfe\mon_pfapp"
C:\Users\LOQ\dev\flutter\bin\cache\dart-sdk\bin\dart.exe analyze
C:\Users\LOQ\dev\flutter\bin\flutter.bat test

cd "C:\Users\LOQ\Documents\yassmin pfe\mon_pfapi"
..\tools\runtime\php-8.4.22\php.exe artisan test''')

    section("12. Debug rapide", [
        "Cette table donne les causes probables et les solutions pour les erreurs les plus courantes pendant la soutenance ou le developpement.",
    ], table=(["Symptome", "Cause probable", "Solution"], [["API ne repond pas", "artisan serve non lance", "Relancer Laravel port 8000"], ["Login impossible telephone", "Mauvaise IP PC", "Utiliser ipconfig et reconstruire/run"], ["ADB vide", "USB debug/RSA absent", "Accepter RSA ou changer cable"], ["dart analyze bloque", "dart.bat wrapper", "Utiliser dart.exe direct"], ["Windows atlstr.h", "ATL absent", "Installer composant Visual Studio"], ["Release ne joint pas HTTP", "Cleartext refuse", "Utiliser HTTPS ou debug"]], [4.2, 5.0, 5.8]))

    section("13. Architecture", [
        "L'architecture separe le frontend Flutter et le backend Laravel. Cette separation permet de tester le frontend en demo et de brancher une API reelle pour valider les donnees persistantes.",
    ], image=diagrams["architecture"])

    section("14. Fichiers importants", [
        "Yassmine doit connaitre ces fichiers pour expliquer ou modifier son application.",
    ], table=(["Fichier", "Role"], [["mon_pfapp/lib/app/mon_pf_app.dart", "Etat global, navigation, panier"], ["mon_pfapp/lib/data/api_client.dart", "HTTP, token, timeout"], ["mon_pfapp/lib/features/client/data/order_service.dart", "Commandes cote app"], ["mon_pfapi/routes/api.php", "Endpoints REST"], ["mon_pfapi/app/Http/Middleware/TokenAuth.php", "Controle token et roles"], ["mon_pfapi/database/seeders/DatabaseSeeder.php", "Donnees de test"]], [7.0, 8.0]))

    section("15. Comptes de test", [
        "Tous les comptes de test utilisent le mot de passe password. Ils sont recrees par migrate:fresh --seed.",
    ], table=(["Email", "Role", "Usage"], [["yassmine@monpf.fr", "client", "Parcours client"], ["admin@monpf.fr", "admin", "Dashboard et stats"], ["serveur@monpf.fr", "serveur", "Flux restaurant"], ["livreur@monpf.fr", "livreur", "Livraison"]], [6.0, 4.0, 5.0]))

    section("16. Scenario de soutenance", [
        "Ce scenario est l'ordre recommande pour presenter l'application sans se perdre.",
    ], numbered=["Presenter la problematique du restaurant.", "Lancer /api/health pour prouver le backend.", "Se connecter client.", "Ajouter des plats et creer une commande.", "Afficher le suivi.", "Se connecter admin.", "Se connecter livreur et montrer la file.", "Montrer les tests automatiques.", "Expliquer les perspectives production."], image=SCREENSHOT_DIR / "contact-sheet.png")

    section("17. Checklist finale", [
        "Avant la soutenance, verifier cette liste dans l'ordre.",
    ], bullets=["Backend Laravel lance sur port 8000.", "Telephone et PC sur meme Wi-Fi si test mobile.", "APK debug installe ou appareil Flutter detecte.", "Comptes de test memorises.", "Rapport PFE et tutoriel PDF disponibles.", "GitHub main a jour.", "Windows ATL installe si build Windows requis."])

    doc = SimpleDocTemplate(str(TUTORIAL_PDF), pagesize=A4, rightMargin=1.7 * cm, leftMargin=1.7 * cm, topMargin=1.65 * cm, bottomMargin=1.55 * cm, title="Tutoriel Mon PF App")
    doc.build(story, onFirstPage=footer, onLaterPages=footer)


def main() -> None:
    PFE_DIR.mkdir(parents=True, exist_ok=True)
    diagrams = build_diagrams()
    build_pfe_docx(diagrams)
    build_tutorial_docx(diagrams)
    build_report_pdf_full(diagrams)
    build_tutorial_pdf_full(diagrams)
    print(f"Generated: {PFE_DOCX}")
    print(f"Generated: {PFE_PDF}")
    print(f"Generated: {TUTORIAL_DOCX}")
    print(f"Generated: {TUTORIAL_PDF}")


if __name__ == "__main__":
    main()
