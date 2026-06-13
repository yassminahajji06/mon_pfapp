from __future__ import annotations

import math
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image as RLImage,
    KeepTogether,
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
TUTORIAL_DIR = ROOT / "docs" / "tutorial"
ASSET_DIR = TUTORIAL_DIR / "assets"
DIAGRAM_DIR = ASSET_DIR / "diagrams"
SCREENSHOT_DIR = ASSET_DIR / "screenshots"
DOCX_PATH = TUTORIAL_DIR / "Mon_PF_App_Tutorial_Yassmine_Hajji.docx"
PDF_PATH = TUTORIAL_DIR / "Mon_PF_App_Tutorial_Yassmine_Hajji.pdf"

ACCENT = RGBColor(229, 57, 53)
BLUE = RGBColor(46, 116, 181)
NAVY = RGBColor(31, 77, 120)
TEXT = RGBColor(31, 41, 51)
MUTED = RGBColor(110, 118, 129)
LIGHT_BLUE = "E8EEF5"
LIGHT_RED = "FDECEC"
LIGHT_GREEN = "EAF7ED"
LIGHT_YELLOW = "FFF7E6"


def font_path(name: str) -> str | None:
    candidates = [
        Path("C:/Windows/Fonts") / name,
        Path("/usr/share/fonts/truetype/dejavu") / name,
    ]
    for candidate in candidates:
        if candidate.exists():
            return str(candidate)
    return None


def pil_font(size: int, bold: bool = False) -> ImageFont.ImageFont:
    if bold:
        for name in ("arialbd.ttf", "DejaVuSans-Bold.ttf"):
            path = font_path(name)
            if path:
                return ImageFont.truetype(path, size)
    for name in ("arial.ttf", "DejaVuSans.ttf"):
        path = font_path(name)
        if path:
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def wrap_text(text: str, font: ImageFont.ImageFont, max_width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        test = f"{current} {word}".strip()
        if font.getbbox(test)[2] <= max_width or not current:
            current = test
        else:
            lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_round_box(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int, int, int],
    title: str,
    lines: Iterable[str],
    *,
    fill: str,
    outline: str = "#D7DBE2",
    title_fill: str = "#1F2933",
    body_fill: str = "#4B5563",
    radius: int = 24,
) -> None:
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=2)
    title_font = pil_font(26, bold=True)
    body_font = pil_font(20)
    draw.text((x1 + 24, y1 + 20), title, fill=title_fill, font=title_font)
    yy = y1 + 64
    for line in lines:
        for wrapped in wrap_text(line, body_font, x2 - x1 - 48):
            draw.text((x1 + 24, yy), wrapped, fill=body_fill, font=body_font)
            yy += 28


def draw_arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    *,
    color: str = "#E53935",
    width: int = 5,
) -> None:
    draw.line([start, end], fill=color, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 18
    points = [
        end,
        (
            int(end[0] - length * math.cos(angle - math.pi / 6)),
            int(end[1] - length * math.sin(angle - math.pi / 6)),
        ),
        (
            int(end[0] - length * math.cos(angle + math.pi / 6)),
            int(end[1] - length * math.sin(angle + math.pi / 6)),
        ),
    ]
    draw.polygon(points, fill=color)


def create_diagram_canvas(title: str) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    image = Image.new("RGB", (1500, 900), "white")
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, 1500, 92), fill="#E53935")
    draw.text((44, 26), title, fill="white", font=pil_font(34, bold=True))
    return image, draw


def build_diagrams() -> dict[str, Path]:
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)

    diagrams: dict[str, Path] = {}

    image, draw = create_diagram_canvas("Architecture globale")
    draw_round_box(
        draw,
        (70, 155, 390, 340),
        "Flutter App",
        ["Ecrans client", "Livreur", "Admin", "Etat UI + panier"],
        fill="#FDECEC",
        outline="#F3B4B4",
    )
    draw_round_box(
        draw,
        (530, 155, 850, 340),
        "AuthService",
        ["HTTP REST", "Timeout 15 s", "Parsing JSON", "Messages d'erreur"],
        fill="#E8EEF5",
        outline="#AFC4DD",
    )
    draw_round_box(
        draw,
        (1010, 155, 1330, 340),
        "API HTTPS",
        ["Laravel attendu", "/login", "/register", "/logout"],
        fill="#EAF7ED",
        outline="#B8DDBE",
    )
    draw_round_box(
        draw,
        (1010, 505, 1330, 690),
        "Base MySQL",
        ["users", "plats", "commandes", "livraisons"],
        fill="#FFF7E6",
        outline="#F0D28A",
    )
    draw_round_box(
        draw,
        (530, 505, 850, 690),
        "Secure Storage",
        ["Token", "Role", "Efface au logout"],
        fill="#F6F7F9",
    )
    draw_round_box(
        draw,
        (70, 505, 390, 690),
        "DemoData",
        ["Jeu local", "Mode demo", "Tests sans backend"],
        fill="#F6F7F9",
    )
    draw_arrow(draw, (390, 245), (530, 245))
    draw_arrow(draw, (850, 245), (1010, 245))
    draw_arrow(draw, (1170, 340), (1170, 505), color="#1E88E5")
    draw_arrow(draw, (690, 340), (690, 505), color="#43A047")
    draw_arrow(draw, (390, 595), (530, 595), color="#7A7F87")
    draw.text((80, 790), "Le meme frontend peut fonctionner en mode demo ou contre une vraie API REST securisee.", fill="#4B5563", font=pil_font(24, bold=True))
    diagrams["architecture"] = DIAGRAM_DIR / "architecture.png"
    image.save(diagrams["architecture"])

    image, draw = create_diagram_canvas("Flux d'authentification")
    boxes = [
        ((60, 180, 310, 340), "1. Formulaire", ["Login ou inscription", "Validation cote UI"], "#FDECEC"),
        ((390, 180, 640, 340), "2. Service", ["AuthService", "POST JSON"], "#E8EEF5"),
        ((720, 180, 970, 340), "3. API", ["Controle identifiants", "Role impose serveur"], "#EAF7ED"),
        ((1050, 180, 1300, 340), "4. Reponse", ["token", "user", "role"], "#FFF7E6"),
    ]
    for xy, title, lines, fill in boxes:
        draw_round_box(draw, xy, title, lines, fill=fill)
    for x in (310, 640, 970):
        draw_arrow(draw, (x, 260), (x + 80, 260))
    draw_round_box(
        draw,
        (390, 510, 640, 690),
        "5. Stockage",
        ["flutter_secure_storage", "token + role"],
        fill="#F6F7F9",
    )
    draw_round_box(
        draw,
        (720, 510, 970, 690),
        "6. Navigation",
        ["Client", "Livreur", "Admin", "selon role"],
        fill="#F6F7F9",
    )
    draw_arrow(draw, (1175, 340), (515, 510), color="#43A047")
    draw_arrow(draw, (640, 600), (720, 600), color="#43A047")
    draw.text((60, 775), "Important: le formulaire public ne choisit jamais le role; le backend retourne le role autorise.", fill="#4B5563", font=pil_font(24, bold=True))
    diagrams["auth_flow"] = DIAGRAM_DIR / "auth_flow.png"
    image.save(diagrams["auth_flow"])

    image, draw = create_diagram_canvas("Cycle de vie d'une commande")
    steps = [
        ((65, 190, 285, 330), "Menu", ["Le client choisit", "des plats"], "#FDECEC"),
        ((355, 190, 575, 330), "Panier", ["Quantites", "Sous-total", "Livraison"], "#E8EEF5"),
        ((645, 190, 865, 330), "Commande", ["Validation", "Etat initial"], "#EAF7ED"),
        ((935, 190, 1155, 330), "Admin", ["Recoit", "Affecte"], "#FFF7E6"),
        ((1225, 190, 1445, 330), "Livreur", ["Accepte", "Met a jour"], "#F6F7F9"),
    ]
    for xy, title, lines, fill in steps:
        draw_round_box(draw, xy, title, lines, fill=fill)
    for x in (285, 575, 865, 1155):
        draw_arrow(draw, (x, 260), (x + 70, 260))
    draw_round_box(
        draw,
        (210, 520, 520, 705),
        "Regle 1",
        ["Une commande ne peut pas etre validee sans plats."],
        fill="#FFF7E6",
        outline="#F0D28A",
    )
    draw_round_box(
        draw,
        (625, 520, 935, 705),
        "Regle 2",
        ["Une commande livree ne peut plus etre modifiee."],
        fill="#FFF7E6",
        outline="#F0D28A",
    )
    draw_round_box(
        draw,
        (1040, 520, 1350, 705),
        "Suivi",
        ["Le client voit la progression et les etapes."],
        fill="#FDECEC",
        outline="#F3B4B4",
    )
    draw_arrow(draw, (1335, 330), (1195, 520), color="#1E88E5")
    diagrams["order_flow"] = DIAGRAM_DIR / "order_flow.png"
    image.save(diagrams["order_flow"])

    image, draw = create_diagram_canvas("Organisation du repository")
    draw_round_box(
        draw,
        (70, 150, 410, 720),
        "Racine",
        ["README.md", ".gitignore", "tools/", "docs/", "mon_pfapp/"],
        fill="#FDECEC",
    )
    draw_round_box(
        draw,
        (545, 150, 885, 365),
        "docs/",
        ["specifications/", "design/", "tutorial/"],
        fill="#E8EEF5",
    )
    draw_round_box(
        draw,
        (545, 505, 885, 720),
        "tools/",
        ["capture_app_screenshots.js", "build_tutorial_pdf.py"],
        fill="#F6F7F9",
    )
    draw_round_box(
        draw,
        (1020, 150, 1400, 720),
        "mon_pfapp/",
        ["lib/app", "lib/core", "lib/domain", "lib/features", "lib/shared", "test/", "android/", "windows/"],
        fill="#EAF7ED",
    )
    draw_arrow(draw, (410, 265), (545, 265))
    draw_arrow(draw, (410, 610), (545, 610), color="#7A7F87")
    draw_arrow(draw, (885, 610), (1020, 610), color="#43A047")
    draw_arrow(draw, (885, 265), (1020, 265))
    diagrams["repo_structure"] = DIAGRAM_DIR / "repo_structure.png"
    image.save(diagrams["repo_structure"])

    image, draw = create_diagram_canvas("Modele de donnees cote Flutter")
    draw_round_box(
        draw,
        (90, 165, 470, 380),
        "UserModel",
        ["id", "nom", "email", "role", "fromJson()"],
        fill="#E8EEF5",
    )
    draw_round_box(
        draw,
        (570, 165, 950, 380),
        "MenuItem",
        ["id, name, category", "description, price", "rating, prepTime", "vegetarian, popular"],
        fill="#FDECEC",
    )
    draw_round_box(
        draw,
        (1030, 165, 1410, 380),
        "CartItem",
        ["item: MenuItem", "quantity", "total", "copyWith()"],
        fill="#EAF7ED",
    )
    draw_round_box(
        draw,
        (570, 535, 950, 750),
        "OrderModel",
        ["id", "client", "address", "status", "time", "amount"],
        fill="#FFF7E6",
    )
    draw_arrow(draw, (1030, 270), (950, 270), color="#43A047")
    draw.text((965, 238), "contient", fill="#4B5563", font=pil_font(20, bold=True))
    draw_arrow(draw, (780, 380), (780, 535), color="#1E88E5")
    draw.text((810, 450), "devient commande", fill="#4B5563", font=pil_font(20, bold=True))
    draw_arrow(draw, (470, 270), (570, 610), color="#7A7F87")
    draw.text((420, 470), "client", fill="#4B5563", font=pil_font(20, bold=True))
    diagrams["data_model"] = DIAGRAM_DIR / "data_model.png"
    image.save(diagrams["data_model"])

    return diagrams


def set_run_font(run, *, size: float | None = None, color: RGBColor | None = None, bold: bool | None = None, italic: bool | None = None, name: str = "Calibri") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, *, bold: bool = False, color: RGBColor | None = None, size: float = 9.5) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, color=color or TEXT, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color: str = "D7DBE2") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(table, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def setup_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, NAVY, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header_p = section.header.paragraphs[0]
    header_p.text = "Mon PF App | Tutoriel technique"
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(header_p.runs[0], size=9, color=MUTED, bold=True)

    footer_p = section.footer.paragraphs[0]
    footer_p.text = "Projet de fin d'etudes - Yassmine Hajji - Application de gestion restaurant"
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(footer_p.runs[0], size=8.5, color=MUTED)


def add_title_page(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = kicker.add_run("GUIDE TECHNIQUE & TUTORIEL")
    set_run_font(r, size=12, color=ACCENT, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Mon PF App")
    set_run_font(r, size=32, color=TEXT, bold=True)
    title.paragraph_format.space_after = Pt(6)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Application mobile de gestion des commandes et livraisons d'un restaurant")
    set_run_font(r, size=14, color=NAVY, bold=True)

    meta = doc.add_table(rows=5, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = False
    set_table_borders(meta)
    set_cell_margins(meta)
    rows = [
        ("Etudiante", "Yassmine Hajji"),
        ("Filiere", "Developpement d'applications"),
        ("Contexte", "Projet de fin d'etudes"),
        ("Version du guide", "13 juin 2026"),
        ("Plateformes ciblees", "Android en priorite, Windows en preparation"),
    ]
    for row, (label, value) in zip(meta.rows, rows):
        row.cells[0].width = Inches(1.8)
        row.cells[1].width = Inches(4.4)
        shade_cell(row.cells[0], LIGHT_BLUE)
        set_cell_text(row.cells[0], label, bold=True, color=NAVY)
        set_cell_text(row.cells[1], value)

    doc.add_paragraph()
    callout = doc.add_table(rows=1, cols=1)
    callout.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(callout, color="F3B4B4")
    set_cell_margins(callout, top=140, bottom=140, start=180, end=180)
    shade_cell(callout.cell(0, 0), LIGHT_RED)
    set_cell_text(
        callout.cell(0, 0),
        "Objectif du document: expliquer clairement ce qui a ete construit, pourquoi ces choix techniques ont ete faits, comment les pieces se connectent, et comment tester/presenter l'application.",
        bold=True,
        color=TEXT,
        size=10.5,
    )
    doc.add_page_break()


def add_paragraph(doc: Document, text: str, *, style: str | None = None, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_run_font(r, bold=True, color=TEXT)
        rest = text[len(bold_lead) :]
        if rest:
            r = p.add_run(rest)
            set_run_font(r, color=TEXT)
    else:
        r = p.add_run(text)
        set_run_font(r, color=TEXT)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run_font(r, color=TEXT)


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_run_font(r, color=TEXT)


def add_matrix(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    set_cell_margins(table)
    hdr = table.rows[0].cells
    for index, header in enumerate(headers):
        if widths:
            hdr[index].width = Inches(widths[index])
        shade_cell(hdr[index], LIGHT_BLUE)
        set_cell_text(hdr[index], header, bold=True, color=NAVY, size=9.3)
    for row_values in rows:
        row = table.add_row().cells
        for index, value in enumerate(row_values):
            if widths:
                row[index].width = Inches(widths[index])
            set_cell_text(row[index], value, size=9.1)
    doc.add_paragraph()


def add_figure(doc: Document, path: Path, caption: str, *, width: float = 6.4) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_run_font(r, size=9, color=MUTED, italic=True)
    cap.paragraph_format.space_after = Pt(8)


def add_code_block(doc: Document, command: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="D7DBE2")
    set_cell_margins(table, top=110, bottom=110, start=150, end=150)
    cell = table.cell(0, 0)
    shade_cell(cell, "F6F7F9")
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(command)
    set_run_font(run, name="Consolas", size=9.2, color=TEXT)
    doc.add_paragraph()


def add_screenshot_grid(doc: Document) -> None:
    screenshots = [
        ("01-login.png", "Connexion"),
        ("02-home.png", "Accueil client"),
        ("03-menu.png", "Menu"),
        ("04-cart.png", "Panier"),
        ("05-tracking.png", "Suivi livraison"),
        ("06-profile.png", "Profil"),
        ("07-driver.png", "Espace livreur"),
        ("08-admin.png", "Administration"),
    ]
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color="E5E7EB")
    set_cell_margins(table, top=80, bottom=80, start=80, end=80)
    for index, (filename, label) in enumerate(screenshots):
        row = index // 2
        col = index % 2
        cell = table.cell(row, col)
        cell.width = Inches(3.1)
        shade_cell(cell, "FFFFFF")
        cell.text = ""
        title = cell.paragraphs[0]
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = title.add_run(label)
        set_run_font(r, size=9.5, color=NAVY, bold=True)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(SCREENSHOT_DIR / filename), width=Inches(2.35))
    doc.add_paragraph()


def build_document(diagrams: dict[str, Path]) -> None:
    TUTORIAL_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_document(doc)
    add_title_page(doc)

    doc.add_heading("1. Vue d'ensemble", level=1)
    add_paragraph(
        doc,
        "Mon PF App est une application Flutter qui digitalise le parcours d'un restaurant: le client consulte le menu, ajoute des plats au panier, valide une commande, suit la livraison, pendant que les espaces livreur et administrateur montrent le pilotage operationnel.",
    )
    add_paragraph(
        doc,
        "Le projet est actuellement livre avec un mode demonstration active par defaut. Ce mode permet de tester l'interface complete sans backend. Quand l'API Laravel/MySQL est prete, l'application bascule vers le vrai serveur avec des variables d'environnement Flutter.",
    )
    add_matrix(
        doc,
        ["Acteur", "Ce qu'il fait dans l'application", "Ecrans concernes"],
        [
            ["Client", "Se connecter, parcourir le menu, gerer le panier, suivre la commande.", "Login, Accueil, Menu, Panier, Suivi, Commandes, Profil"],
            ["Livreur", "Voir les livraisons disponibles, accepter/refuser et suivre les indicateurs.", "Tableau de bord livreur"],
            ["Administrateur", "Observer les ventes, commandes, clients actifs, livreurs et alertes stock.", "Tableau de bord administration"],
            ["Serveur", "Representable dans le flux operationnel/admin; peut devenir un espace dedie si le jury demande une separation.", "Extension future"],
        ],
        widths=[1.3, 3.3, 1.9],
    )
    add_figure(doc, diagrams["architecture"], "Figure 1 - Architecture cible et mode demo.", width=6.5)

    doc.add_heading("2. Stack technique et raisons des choix", level=1)
    add_matrix(
        doc,
        ["Technologie", "Ou elle se trouve", "Pourquoi ce choix"],
        [
            ["Flutter / Dart", "mon_pfapp/lib", "Un seul codebase pour Android, Windows, web de test, et plus tard iOS/Linux/macOS si les toolchains sont installees."],
            ["Material 3", "MyApp + shared/widgets/app_ui.dart", "Composants coherents, themes faciles, rendu propre sur mobile."],
            ["HTTP REST", "features/auth/data/auth_service.dart", "Integration simple avec un backend Laravel classique."],
            ["flutter_secure_storage", "core/storage.dart", "Stocker le token hors SharedPreferences pour reduire le risque de fuite."],
            ["Laravel + MySQL attendu", "Backend externe", "Correspond au cahier de charge: utilisateurs, menu, commandes, livraisons, statistiques."],
            ["Gradle / Android SDK", "android/", "Build debug et release APK pour test et livraison Android."],
        ],
        widths=[1.45, 2.2, 2.85],
    )
    add_paragraph(
        doc,
        "Choix important: le frontend est pret pour une API reelle, mais le mode demo reste indispensable pour une soutenance ou une production de presentation, car il evite qu'un probleme serveur bloque la demonstration.",
        bold_lead="Choix important:",
    )

    doc.add_heading("3. Organisation du repository", level=1)
    add_figure(doc, diagrams["repo_structure"], "Figure 2 - Structure professionnelle du depot.", width=6.5)
    add_matrix(
        doc,
        ["Dossier", "Role"],
        [
            ["docs/specifications", "Documents originaux fournis: cahier de charge et analyse/conception."],
            ["docs/design", "Reference visuelle de la maquette utilisee pour guider l'UI."],
            ["docs/tutorial", "Ce guide, les diagrammes et les captures d'ecran."],
            ["tools", "Scripts reproductibles pour capturer l'app et reconstruire le guide."],
            ["mon_pfapp/lib/app", "Application shell, navigation et etat global de demonstration."],
            ["mon_pfapp/lib/core", "Configuration API, stockage securise, validateurs."],
            ["mon_pfapp/lib/domain/models", "Objets metier: user, menu item, cart item, order."],
            ["mon_pfapp/lib/features", "Ecrans et logique par domaine: auth, client, operations."],
            ["mon_pfapp/lib/shared", "Widgets reutilisables, couleurs, cartes, navigation basse."],
        ],
        widths=[2.2, 4.3],
    )

    doc.add_heading("4. Architecture applicative Flutter", level=1)
    add_paragraph(
        doc,
        "Le point d'entree est main.dart. Il lance MyApp, qui configure MaterialApp, le theme et la page principale MonPfApp. MonPfApp est un StatefulWidget qui possede l'etat simple de demonstration: ecran courant, utilisateur connecte et panier.",
    )
    add_numbered(
        doc,
        [
            "L'utilisateur arrive sur LoginScreen.",
            "LoginScreen appelle AuthService.login ou RegisterScreen appelle AuthService.register.",
            "Si l'authentification reussit, MonPfApp stocke le UserModel et bascule vers HomeScreen.",
            "La navigation interne change la valeur _screen: home, menu, cart, tracking, orders, profile, driver ou admin.",
            "Les actions panier modifient la liste CartItem dans MonPfApp, puis les ecrans se reconstruisent automatiquement.",
        ],
    )
    add_matrix(
        doc,
        ["Element", "Responsabilite", "Connexion avec le reste"],
        [
            ["MonPfApp", "Etat d'application et navigation simple.", "Transmet callbacks et donnees aux ecrans."],
            ["AuthService", "Inscription, connexion, deconnexion et erreurs API.", "Utilise AppConstants, Storage et UserModel."],
            ["Storage", "Token et role dans secure storage.", "Utilise par AuthService apres reponse API."],
            ["DemoData", "Donnees locales pour menus, commandes et utilisateur.", "Permet de tester sans backend."],
            ["AppBottomNav", "Navigation client persistante.", "Declenche onNavigate fourni par MonPfApp."],
        ],
        widths=[1.45, 2.65, 2.4],
    )

    doc.add_heading("5. Authentification et securite", level=1)
    add_figure(doc, diagrams["auth_flow"], "Figure 3 - Connexion, role et stockage du token.", width=6.5)
    add_paragraph(
        doc,
        "Avant les corrections, l'application pouvait etre trop dependante d'une adresse locale, stocker des donnees sensibles trop simplement, et permettre une selection de role cote inscription. Ces points ont ete corriges pour rapprocher le projet d'une logique production.",
    )
    add_matrix(
        doc,
        ["Risque initial", "Solution appliquee", "Impact"],
        [
            ["API en clair ou IP locale fixe", "API_BASE_URL en dart-define avec valeur HTTPS par defaut.", "Deploiement flexible et plus securise."],
            ["Role choisi par le public", "RegisterScreen cree toujours un client; le role vient du serveur.", "Evite l'auto-promotion en admin/livreur."],
            ["Token en stockage simple", "flutter_secure_storage.", "Meilleure protection locale du jeton."],
            ["Serveur lent ou indisponible", "Timeout 15 secondes et messages d'erreur propres.", "L'UI ne reste pas bloquee sans explication."],
            ["Retour asynchrone apres fermeture d'ecran", "Checks mounted dans les ecrans auth/profil.", "Evite les setState invalides."],
        ],
        widths=[1.75, 2.55, 2.2],
    )
    add_paragraph(
        doc,
        "En mode reel, l'API doit retourner un objet user avec un role. Le frontend lit ce role, mais la decision de securite doit rester cote backend: seul le serveur sait si un utilisateur est client, livreur ou administrateur.",
    )

    doc.add_heading("6. Modele de donnees", level=1)
    add_figure(doc, diagrams["data_model"], "Figure 4 - Classes metier principales dans Flutter.", width=6.5)
    add_matrix(
        doc,
        ["Classe", "Champs importants", "Utilisation"],
        [
            ["UserModel", "id, nom, email, role", "Representer l'utilisateur authentifie et personnaliser la navigation."],
            ["MenuItem", "id, name, category, description, price, rating, prepTime", "Afficher les plats, filtrer par categorie et ajouter au panier."],
            ["CartItem", "item, quantity, total", "Lier un plat a une quantite et calculer le total ligne."],
            ["OrderModel", "id, client, address, status, time, amount", "Afficher les commandes dans les espaces operations."],
        ],
        widths=[1.35, 2.6, 2.55],
    )

    doc.add_heading("7. Algorithmes et regles metier", level=1)
    add_paragraph(
        doc,
        "Le projet n'utilise pas un algorithme complexe de type IA. Il utilise plutot des algorithmes applicatifs simples, lisibles et defendables: gestion de panier, calcul total, routage d'ecran, controle role/token et progression commande.",
    )
    add_matrix(
        doc,
        ["Mecanisme", "Comment ca marche", "Pourquoi c'est important"],
        [
            ["Ajout panier", "Cherche l'item par id. S'il n'existe pas, ajoute CartItem quantite 1. Sinon reconstruit la liste avec quantite + 1.", "Evite les doublons et garde un etat immutable simple."],
            ["Quantite", "Applique delta +1 ou -1 puis clamp entre 1 et 99.", "Empêche les quantites zero/negatives et les valeurs absurdes."],
            ["Total panier", "Chaque CartItem expose total = prix * quantite; l'ecran panier additionne sous-total, livraison et remise.", "Calcul clair pour la facture client."],
            ["Navigation", "Un switch sur _screen choisit quel ecran afficher.", "Simple a presenter, suffisant pour une demo de projet."],
            ["Auth API", "POST JSON, decode JSON, controle status code, token + user requis, timeout et erreurs fallback.", "Robustesse reseau et messages comprehensibles."],
            ["Role", "Inscription publique force client; roles sensibles viennent du backend.", "Regle centrale de securite."],
            ["Cycle commande", "Commande confirmee, preparation, livreur en route, livraison effectuee.", "Correspond au suivi attendu dans le cahier de charge."],
        ],
        widths=[1.45, 3.05, 2.0],
    )
    add_figure(doc, diagrams["order_flow"], "Figure 5 - Regles de validation et suivi d'une commande.", width=6.5)

    doc.add_heading("8. Parcours utilisateur et captures", level=1)
    add_paragraph(
        doc,
        "Les captures ci-dessous viennent de la version web release de l'application Flutter. Elles servent a expliquer le parcours et a verifier que l'interface est coherente avec la maquette fournie.",
    )
    add_screenshot_grid(doc)

    doc.add_heading("9. Tests, builds et etat actuel", level=1)
    add_matrix(
        doc,
        ["Controle", "Commande / resultat", "Etat"],
        [
            ["Analyse Dart", "dart analyze", "OK"],
            ["Tests Flutter", "flutter test", "5 tests OK"],
            ["Web release", "flutter build web --release --dart-define=DEMO_MODE=true", "OK"],
            ["Android debug", "flutter build apk --debug", "OK"],
            ["Android release signe", "flutter build apk --release", "OK, APK genere"],
            ["Windows release", "flutter build windows --release", "Bloque par composant ATL manquant dans Visual Studio Build Tools."],
        ],
        widths=[1.45, 3.25, 1.8],
    )
    add_paragraph(
        doc,
        "Point Windows: l'erreur actuelle vient de flutter_secure_storage_windows qui a besoin de atlstr.h. Il faut installer le composant Visual Studio 'C++ ATL for latest v142 build tools'. Apres cela, relancer flutter build windows --release.",
        bold_lead="Point Windows:",
    )

    doc.add_heading("10. Comment lancer et tester", level=1)
    add_paragraph(doc, "Depuis la racine du projet:")
    add_code_block(doc, "cd mon_pfapp\nflutter pub get\nflutter test\nflutter run")
    add_paragraph(doc, "Pour tester en mode demo avec Chrome sur le port utilise pour les captures:")
    add_code_block(doc, "flutter run -d chrome --web-port 57321")
    add_paragraph(doc, "Pour brancher une vraie API HTTPS:")
    add_code_block(doc, "flutter run --dart-define=DEMO_MODE=false --dart-define=API_BASE_URL=https://votre-domaine/api")
    add_paragraph(doc, "Pour installer l'APK release sur un telephone Android connecte:")
    add_code_block(doc, "adb install -r build/app/outputs/flutter-apk/app-release.apk")

    doc.add_heading("11. Ce qui reste pour une vraie production", level=1)
    add_bullets(
        doc,
        [
            "Brancher les routes Laravel/MySQL finales: register, login, logout, menu, commandes, affectation livreur, suivi.",
            "Ajouter une gestion serveur des roles avec middleware d'autorisation.",
            "Remplacer certaines donnees DemoData par des appels API pagines/cachees.",
            "Ajouter notifications temps reel ou polling pour le suivi de livraison.",
            "Garder les cles de signature Android hors Git et utiliser une strategie de sauvegarde secrete.",
            "Installer le composant ATL pour produire la version Windows.",
            "Ajouter CI/CD GitHub Actions pour analyse, tests et builds automatiques.",
        ],
    )

    doc.add_heading("12. Comment presenter le projet au jury", level=1)
    add_numbered(
        doc,
        [
            "Commencer par la problematique: erreurs, retards et manque de suivi dans une gestion restaurant manuelle.",
            "Presenter la solution: une app Flutter qui connecte client, restaurant, administrateur et livreur.",
            "Montrer le parcours client complet: login, menu, panier, validation, suivi.",
            "Montrer les espaces operations: livreur puis administration.",
            "Expliquer les choix securite: HTTPS/env config, secure storage, roles serveur, timeouts.",
            "Expliquer la structure professionnelle du repo et les tests.",
            "Finir par les limites honnetes: backend reel, temps reel, paiement et build Windows apres ATL.",
        ],
    )
    add_matrix(
        doc,
        ["Question probable", "Reponse courte"],
        [
            ["Pourquoi Flutter ?", "Pour livrer rapidement Android et garder une base compatible web/desktop."],
            ["Pourquoi secure storage ?", "Parce qu'un token d'authentification ne doit pas rester dans un stockage simple."],
            ["Pourquoi supprimer le choix du role ?", "Un utilisateur public ne doit jamais pouvoir se creer admin ou livreur."],
            ["Pourquoi le mode demo ?", "Pour garantir une soutenance testable meme si le backend n'est pas encore branche."],
            ["Pourquoi Windows est bloque ?", "Il manque le composant ATL de Visual Studio, requis par le plugin secure storage Windows."],
        ],
        widths=[2.1, 4.4],
    )

    doc.add_heading("Annexe A - Fichiers importants", level=1)
    add_matrix(
        doc,
        ["Fichier", "A quoi il sert"],
        [
            ["lib/app/mon_pf_app.dart", "Shell applicatif, navigation et panier de demonstration."],
            ["lib/core/constants.dart", "Nom app, API_BASE_URL, DEMO_MODE, timeout."],
            ["lib/core/storage.dart", "Sauvegarde securisee token/role."],
            ["lib/features/auth/data/auth_service.dart", "Inscription, connexion, deconnexion, erreurs API."],
            ["lib/features/client/presentation", "Ecrans client: accueil, menu, panier, suivi, profil."],
            ["lib/features/operations/presentation", "Ecrans livreur et administration."],
            ["mon_pfapp/docs", "Documentation technique courte par sujet."],
            ["docs/tutorial", "Guide complet, images, diagrammes et PDF."],
        ],
        widths=[2.35, 4.15],
    )

    doc.save(DOCX_PATH)


def pdf_styles() -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "title",
            parent=base["Title"],
            fontName="Helvetica-Bold",
            fontSize=30,
            leading=34,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#1F2933"),
            spaceAfter=10,
        ),
        "subtitle": ParagraphStyle(
            "subtitle",
            parent=base["Normal"],
            fontName="Helvetica-Bold",
            fontSize=13,
            leading=17,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#1F4D78"),
            spaceAfter=18,
        ),
        "h1": ParagraphStyle(
            "h1",
            parent=base["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=16,
            leading=20,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=14,
            spaceAfter=8,
        ),
        "h2": ParagraphStyle(
            "h2",
            parent=base["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=12.5,
            leading=16,
            textColor=colors.HexColor("#1F4D78"),
            spaceBefore=10,
            spaceAfter=6,
        ),
        "body": ParagraphStyle(
            "body",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=9.8,
            leading=13.2,
            textColor=colors.HexColor("#1F2933"),
            alignment=TA_LEFT,
            spaceAfter=6,
        ),
        "small": ParagraphStyle(
            "small",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=8.6,
            leading=11,
            textColor=colors.HexColor("#4B5563"),
            spaceAfter=4,
        ),
        "caption": ParagraphStyle(
            "caption",
            parent=base["Italic"],
            fontName="Helvetica-Oblique",
            fontSize=8.2,
            leading=10,
            textColor=colors.HexColor("#6B7280"),
            alignment=TA_CENTER,
            spaceAfter=7,
        ),
        "callout": ParagraphStyle(
            "callout",
            parent=base["BodyText"],
            fontName="Helvetica-Bold",
            fontSize=9.4,
            leading=12.5,
            textColor=colors.HexColor("#1F2933"),
            spaceAfter=0,
        ),
        "code": ParagraphStyle(
            "code",
            parent=base["Code"],
            fontName="Courier",
            fontSize=8.2,
            leading=10.5,
            textColor=colors.HexColor("#111827"),
        ),
    }


def p(text: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(text.replace("\n", "<br/>"), style)


def pdf_table(headers: list[str], rows: list[list[str]], styles: dict[str, ParagraphStyle], widths: list[float]) -> Table:
    data = [[p(f"<b>{header}</b>", styles["small"]) for header in headers]]
    data.extend([[p(cell, styles["small"]) for cell in row] for row in rows])
    table = Table(data, colWidths=[width * inch for width in widths], hAlign="CENTER")
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#1F4D78")),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#D7DBE2")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    return table


def pdf_image(path: Path, width: float) -> RLImage:
    with Image.open(path) as img:
        aspect = img.height / img.width
    return RLImage(str(path), width=width * inch, height=width * aspect * inch)


def pdf_callout(text: str, styles: dict[str, ParagraphStyle], fill: str = "#FDECEC") -> Table:
    table = Table([[p(text, styles["callout"])]], colWidths=[6.35 * inch], hAlign="CENTER")
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(fill)),
                ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor("#F3B4B4")),
                ("LEFTPADDING", (0, 0), (-1, -1), 10),
                ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
            ]
        )
    )
    return table


def pdf_bullets(items: Iterable[str], styles: dict[str, ParagraphStyle], numbered: bool = False) -> Table:
    rows = []
    for index, item in enumerate(items, start=1):
        marker = f"{index}." if numbered else "•"
        rows.append([p(marker, styles["body"]), p(item, styles["body"])])
    table = Table(rows, colWidths=[0.28 * inch, 6.05 * inch], hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ALIGN", (0, 0), (0, -1), "RIGHT"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (0, -1), 6),
                ("RIGHTPADDING", (1, 0), (1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 1),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
            ]
        )
    )
    return table


def pdf_code(command: str, styles: dict[str, ParagraphStyle]) -> Table:
    table = Table([[p(command, styles["code"])]], colWidths=[6.35 * inch], hAlign="CENTER")
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F6F7F9")),
                ("BOX", (0, 0), (-1, -1), 0.4, colors.HexColor("#D7DBE2")),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ]
        )
    )
    return table


def add_pdf_footer(canvas, doc) -> None:
    canvas.saveState()
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(colors.HexColor("#6B7280"))
    canvas.drawString(inch, 0.45 * inch, "Mon PF App - Tutoriel technique")
    canvas.drawRightString(7.5 * inch, 0.45 * inch, f"Page {doc.page}")
    canvas.restoreState()


def build_pdf(diagrams: dict[str, Path]) -> None:
    styles = pdf_styles()
    story = []

    story.extend(
        [
            Spacer(1, 1.0 * inch),
            p("GUIDE TECHNIQUE & TUTORIEL", styles["subtitle"]),
            p("Mon PF App", styles["title"]),
            p("Application mobile de gestion des commandes et livraisons d'un restaurant", styles["subtitle"]),
            Spacer(1, 0.15 * inch),
            pdf_table(
                ["Champ", "Valeur"],
                [
                    ["Etudiante", "Yassmine Hajji"],
                    ["Filiere", "Developpement d'applications"],
                    ["Contexte", "Projet de fin d'etudes"],
                    ["Version du guide", "13 juin 2026"],
                    ["Plateformes ciblees", "Android en priorite, Windows en preparation"],
                ],
                styles,
                [1.55, 4.75],
            ),
            Spacer(1, 0.25 * inch),
            pdf_callout(
                "Objectif: expliquer ce qui a ete construit, pourquoi les choix techniques ont ete faits, comment les pieces se connectent, et comment tester/presenter l'application.",
                styles,
            ),
            PageBreak(),
        ]
    )

    story.append(p("1. Vue d'ensemble", styles["h1"]))
    story.append(
        p(
            "Mon PF App digitalise le parcours d'un restaurant: le client consulte le menu, ajoute des plats au panier, valide une commande et suit la livraison. Les espaces livreur et administrateur montrent le pilotage operationnel.",
            styles["body"],
        )
    )
    story.append(
        p(
            "Le mode demonstration est actif par defaut pour tester toute l'interface sans backend. Quand l'API Laravel/MySQL sera prete, l'application pourra basculer vers le vrai serveur avec les dart-define Flutter.",
            styles["body"],
        )
    )
    story.append(
        pdf_table(
            ["Acteur", "Responsabilite", "Ecrans"],
            [
                ["Client", "Se connecter, parcourir le menu, gerer le panier, suivre la commande.", "Login, Accueil, Menu, Panier, Suivi, Profil"],
                ["Livreur", "Voir les livraisons disponibles, accepter/refuser, suivre ses indicateurs.", "Tableau de bord livreur"],
                ["Administrateur", "Observer ventes, commandes, clients, livreurs et alertes stock.", "Tableau de bord administration"],
                ["Serveur", "Representable dans le flux operationnel; peut devenir un espace dedie.", "Extension future"],
            ],
            styles,
            [1.05, 3.05, 2.2],
        )
    )
    story.append(pdf_image(diagrams["architecture"], 6.45))
    story.append(p("Figure 1 - Architecture cible et mode demo.", styles["caption"]))

    story.append(p("2. Stack technique", styles["h1"]))
    story.append(
        pdf_table(
            ["Technologie", "Ou", "Pourquoi"],
            [
                ["Flutter / Dart", "mon_pfapp/lib", "Un codebase pour Android, Windows, web de test et futures plateformes."],
                ["Material 3", "MyApp + shared/widgets", "Composants propres et theme coherent."],
                ["HTTP REST", "AuthService", "Integration simple avec Laravel."],
                ["flutter_secure_storage", "Storage", "Protection locale du token et du role."],
                ["Laravel + MySQL", "Backend attendu", "Modele classique pour users, plats, commandes, livraisons."],
                ["Gradle / Android SDK", "android/", "Build APK debug/release."],
            ],
            styles,
            [1.35, 1.8, 3.15],
        )
    )

    story.append(p("3. Repository et architecture Flutter", styles["h1"]))
    story.append(pdf_image(diagrams["repo_structure"], 6.45))
    story.append(p("Figure 2 - Organisation professionnelle du depot.", styles["caption"]))
    story.append(
        pdf_table(
            ["Zone", "Role"],
            [
                ["docs/specifications", "Cahier de charge et analyse/conception originaux."],
                ["docs/design", "Reference visuelle de la maquette."],
                ["docs/tutorial", "PDF, DOCX, diagrammes et captures."],
                ["tools", "Scripts de capture et generation du guide."],
                ["lib/app", "Shell applicatif, navigation et panier de demonstration."],
                ["lib/core", "Configuration API, secure storage, validateurs."],
                ["lib/features", "Ecrans par domaine: auth, client, operations."],
            ],
            styles,
            [1.75, 4.55],
        )
    )
    story.append(
        p(
            "MonPfApp possede l'etat de demonstration: ecran courant, utilisateur connecte et panier. Les ecrans recoivent des callbacks, ce qui garde la navigation lisible pour la soutenance.",
            styles["body"],
        )
    )

    story.append(p("4. Authentification et securite", styles["h1"]))
    story.append(pdf_image(diagrams["auth_flow"], 6.45))
    story.append(p("Figure 3 - Connexion, role et stockage securise.", styles["caption"]))
    story.append(
        pdf_table(
            ["Risque initial", "Correction", "Impact"],
            [
                ["API locale/en clair", "API_BASE_URL par dart-define et HTTPS par defaut.", "Configuration propre pour prod."],
                ["Role choisi publiquement", "Inscription publique toujours client; roles serveur.", "Evite l'auto-promotion."],
                ["Token stockage simple", "flutter_secure_storage.", "Jeton mieux protege."],
                ["Serveur lent", "Timeout 15 s + messages d'erreur.", "UI plus robuste."],
                ["Async apres fermeture", "Checks mounted.", "Evite les setState invalides."],
            ],
            styles,
            [1.55, 2.55, 2.2],
        )
    )

    story.append(p("5. Modele de donnees", styles["h1"]))
    story.append(pdf_image(diagrams["data_model"], 6.45))
    story.append(p("Figure 4 - Classes metier cote Flutter.", styles["caption"]))
    story.append(
        pdf_table(
            ["Classe", "Champs", "Utilisation"],
            [
                ["UserModel", "id, nom, email, role", "Utilisateur authentifie et navigation selon role."],
                ["MenuItem", "id, name, category, price, rating, prepTime", "Afficher plats, filtrer et ajouter au panier."],
                ["CartItem", "item, quantity, total", "Calculer le total ligne et gerer quantites."],
                ["OrderModel", "id, client, address, status, time, amount", "Afficher les commandes operations."],
            ],
            styles,
            [1.15, 2.45, 2.7],
        )
    )

    story.append(p("6. Algorithmes et regles metier", styles["h1"]))
    story.append(
        p(
            "L'application utilise des algorithmes applicatifs simples et defendables: panier, total, routage d'ecran, controle token/role, decode JSON, timeout et progression commande.",
            styles["body"],
        )
    )
    story.append(
        pdf_table(
            ["Mecanisme", "Comment ca marche", "Pourquoi"],
            [
                ["Ajout panier", "Recherche par id; ajoute un CartItem ou incremente la quantite.", "Evite les doublons."],
                ["Quantite", "Delta +1/-1 puis clamp entre 1 et 99.", "Pas de quantite negative."],
                ["Total", "CartItem.total = prix * quantite; panier additionne sous-total/livraison/remise.", "Facture claire."],
                ["Navigation", "switch sur _screen.", "Simple a expliquer."],
                ["Auth API", "POST JSON, controle status, token+user requis, timeout.", "Robustesse reseau."],
                ["Role", "Client a l'inscription, role sensible via backend.", "Securite."],
            ],
            styles,
            [1.25, 3.25, 1.8],
        )
    )
    story.append(pdf_image(diagrams["order_flow"], 6.45))
    story.append(p("Figure 5 - Cycle de vie et regles d'une commande.", styles["caption"]))

    story.append(p("7. Captures de l'application", styles["h1"]))
    story.append(
        p(
            "Les captures viennent de la version web release Flutter. Elles montrent que le parcours est testable meme sans backend.",
            styles["body"],
        )
    )
    story.append(pdf_image(SCREENSHOT_DIR / "contact-sheet.png", 6.45))
    story.append(p("Figure 6 - Vue rapide des ecrans principaux.", styles["caption"]))

    screen_rows = []
    screens = [
        ("01-login.png", "Connexion"),
        ("02-home.png", "Accueil"),
        ("03-menu.png", "Menu"),
        ("04-cart.png", "Panier"),
        ("05-tracking.png", "Suivi"),
        ("07-driver.png", "Livreur"),
        ("08-admin.png", "Admin"),
        ("06-profile.png", "Profil"),
    ]
    for left, right in zip(screens[0::2], screens[1::2]):
        screen_rows.append(
            [
                [p(f"<b>{left[1]}</b>", styles["small"]), pdf_image(SCREENSHOT_DIR / left[0], 1.7)],
                [p(f"<b>{right[1]}</b>", styles["small"]), pdf_image(SCREENSHOT_DIR / right[0], 1.7)],
            ]
        )
    screen_table = Table(screen_rows, colWidths=[3.1 * inch, 3.1 * inch], hAlign="CENTER")
    screen_table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#E5E7EB")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    story.append(screen_table)

    story.append(PageBreak())
    story.append(p("8. Tests, builds et production", styles["h1"]))
    story.append(
        pdf_table(
            ["Controle", "Commande / resultat", "Etat"],
            [
                ["Analyse Dart", "dart analyze", "OK"],
                ["Tests Flutter", "flutter test", "5 tests OK"],
                ["Web release", "flutter build web --release --dart-define=DEMO_MODE=true", "OK"],
                ["Android debug", "flutter build apk --debug", "OK"],
                ["Android release signe", "flutter build apk --release", "OK, APK genere"],
                ["Windows release", "flutter build windows --release", "Bloque par ATL manquant"],
            ],
            styles,
            [1.45, 3.25, 1.6],
        )
    )
    story.append(
        pdf_callout(
            "Windows: installer le composant Visual Studio 'C++ ATL for latest v142 build tools' pour resoudre l'erreur atlstr.h de flutter_secure_storage_windows.",
            styles,
            "#FFF7E6",
        )
    )
    story.append(p("Commandes utiles", styles["h2"]))
    story.append(pdf_code("cd mon_pfapp\nflutter pub get\nflutter test\nflutter run", styles))
    story.append(pdf_code("flutter run -d chrome --web-port 57321", styles))
    story.append(pdf_code("flutter run --dart-define=DEMO_MODE=false --dart-define=API_BASE_URL=https://votre-domaine/api", styles))
    story.append(pdf_code("adb install -r build/app/outputs/flutter-apk/app-release.apk", styles))

    story.append(p("9. Reste a faire pour une vraie production", styles["h1"]))
    story.append(
        pdf_bullets(
            [
                "Brancher les routes Laravel/MySQL finales: auth, menu, commandes, affectation livreur, suivi.",
                "Ajouter middleware backend pour les roles et permissions.",
                "Remplacer progressivement DemoData par des appels API.",
                "Ajouter notification temps reel ou polling pour le suivi livraison.",
                "Garder les cles de signature Android hors Git.",
                "Installer ATL pour la version Windows.",
                "Ajouter CI/CD GitHub Actions pour analyse, tests et builds.",
            ],
            styles,
        )
    )

    story.append(PageBreak())
    story.append(p("10. Script de presentation au jury", styles["h1"]))
    story.append(
        pdf_bullets(
            [
                "Presenter la problematique: erreurs, retards et manque de visibilite dans une gestion manuelle.",
                "Presenter la solution: app Flutter connectant client, restaurant, livreur et admin.",
                "Demontrer le parcours client: login, menu, panier, suivi.",
                "Demontrer les espaces operations: livreur puis admin.",
                "Expliquer les choix securite: HTTPS/env, secure storage, roles serveur, timeouts.",
                "Conclure sur les tests, l'APK Android et les limites honnetes.",
            ],
            styles,
            numbered=True,
        )
    )
    story.append(Spacer(1, 0.18 * inch))
    story.append(p("Questions probables du jury", styles["h2"]))
    story.append(
        pdf_table(
            ["Question", "Reponse courte"],
            [
                ["Pourquoi Flutter ?", "Pour livrer rapidement Android et garder une base compatible web/desktop."],
                ["Pourquoi secure storage ?", "Un token d'authentification ne doit pas rester dans un stockage simple."],
                ["Pourquoi supprimer le choix du role ?", "Un utilisateur public ne doit jamais pouvoir se creer admin ou livreur."],
                ["Pourquoi le mode demo ?", "Pour garantir une soutenance testable meme si le backend n'est pas encore branche."],
                ["Pourquoi Windows est bloque ?", "Il manque le composant ATL de Visual Studio, requis par le plugin secure storage Windows."],
            ],
            styles,
            [2.05, 4.25],
        )
    )

    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        rightMargin=0.8 * inch,
        leftMargin=0.8 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
        title="Mon PF App - Tutoriel technique",
        author="Yassmine Hajji",
    )
    doc.build(story, onFirstPage=add_pdf_footer, onLaterPages=add_pdf_footer)


def main() -> None:
    diagrams = build_diagrams()
    build_document(diagrams)
    build_pdf(diagrams)
    print(DOCX_PATH)
    print(PDF_PATH)


if __name__ == "__main__":
    main()
